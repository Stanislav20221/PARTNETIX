from django.contrib.auth.decorators import login_required
from django.shortcuts import render
from django.contrib.auth.views import LoginView
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from datetime import datetime
from django.contrib import messages
from django.shortcuts import redirect
from django import forms
from django.conf import settings
from django.contrib.auth import authenticate, get_user_model
from django.contrib.auth.forms import AuthenticationForm
from offers.models import Offer, PartnerLink
from django.shortcuts import redirect, get_object_or_404
import random
from offers.models import Offer, PartnerLink, OfferVisit, PartnerRegistration, Lead
from django.db.models import Sum, Q, Count
from users.models import Accrual, WithdrawalRequest, BankBin, PartnerPaymentDetails, Tariff, EducationVideo, EducationDocument, MarketingMaterial
from decimal import Decimal
from django.db import transaction
import requests
from django.http import JsonResponse
from django.templatetags.static import static
from cryptography.fernet import Fernet
from functools import wraps
from .models import CompanyNews
from users.models import ChatTopic, ChatMessage, User
from .telegram import sync_telegram_news
from datetime import timedelta
from django.utils import timezone
from datetime import timedelta
from django.core.paginator import Paginator
import secrets
import string
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT 
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from django.http import HttpResponse
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import csv
def partner_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('partner_login')

        if request.user.role != 'partner':
            return redirect('/users/dashboard/')

        return view_func(request, *args, **kwargs)

    return wrapper
class PartnerEmailLoginForm(AuthenticationForm):
    username = forms.EmailField(
        label='Email',
        widget=forms.EmailInput(attrs={
            'autofocus': True,
            'placeholder': 'Email',
        })
    )

    def clean(self):
        email = self.cleaned_data.get('username')
        password = self.cleaned_data.get('password')

        if email and password:
            UserModel = get_user_model()

            try:
                user = UserModel.objects.get(email=email)
            except UserModel.DoesNotExist:
                raise forms.ValidationError('Неверный email или пароль.')

            self.user_cache = authenticate(
                self.request,
                username=user.username,
                password=password
            )

            if self.user_cache is None:
                raise forms.ValidationError('Неверный email или пароль.')

        return self.cleaned_data
        
class PartnerLoginView(LoginView):
    template_name = 'partner_portal/login.html'
    authentication_form = PartnerEmailLoginForm

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated:
            return redirect('partner_dashboard')
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        user = form.get_user()

        if not hasattr(user, 'partner_registration'):
            messages.error(self.request, 'Доступ только для партнеров.')
            return self.form_invalid(form)

        if not user.partner_registration.is_approved:
            messages.error(self.request, 'Ваш аккаунт партнера еще не подтвержден.')
            return self.form_invalid(form)

        return super().form_valid(form)

    def get_success_url(self):
        return reverse_lazy('partner_dashboard')
    
@login_required
@partner_required
def dashboard(request):
    accruals_list = Accrual.objects.filter(
    partner=request.user
    )

    balance = accruals_list.filter(
        payout_status='available'
    ).aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    total_accrued = accruals_list.aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    total_leads = Lead.objects.filter(
        partner=request.user
    ).count()

    total_deals = Lead.objects.filter(
        partner=request.user,
        status='deal'
    ).count()
    new_leads = Lead.objects.filter(
        partner=request.user,
        status='new'
    ).count()

    in_work_leads = Lead.objects.filter(
        partner=request.user,
        status='in_work'
    ).count()

    cancelled_leads = Lead.objects.filter(
        partner=request.user,
        status='cancelled'
    ).count()
    referral_links = PartnerLink.objects.filter(
        partner=request.user
    ).select_related('offer').order_by('-created_at')
    news_list = CompanyNews.objects.all()[:5]
    for link in referral_links:
        link.visits_count = link.visits.count()

        link.registrations_count = link.leads.count()

        link.deals_count = link.leads.filter(
            status='deal'
        ).count()

        link.deals_amount = (
            link.leads.filter(status='deal')
            .aggregate(total=Sum('deal_amount'))['total'] or 0
        )
    latest_news_id = news_list[0].id if news_list else None
    return render(request, 'partner_portal/dashboard.html', {
        'active_page': 'dashboard',
        'balance': balance,
        'total_accrued': total_accrued,
        'total_leads': total_leads,
        'total_deals': total_deals,
        'new_leads': new_leads,
        'in_work_leads': in_work_leads,
        'cancelled_leads': cancelled_leads,
        'funnel_data': [
            total_leads,
            new_leads,
            in_work_leads,
            cancelled_leads,
            total_deals,
        ],
        'referral_links': referral_links,
        'news_list': news_list,
        'latest_news_id': latest_news_id,
    })
@login_required
@partner_required
def accruals(request):
    accruals_list = Accrual.objects.filter(
        partner=request.user
    ).select_related(
        'lead',
        'offer',
        'admin'
    ).order_by('-created_at')
    new_accruals_list = accruals_list.filter(
        payout_status='available'
    )
    history_accruals_list = accruals_list.exclude(
        payout_status='available'
    )
    withdrawal_requests = WithdrawalRequest.objects.filter(
        partner=request.user
    ).order_by('-created_at')

    total_accrued = accruals_list.aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    paid_amount = accruals_list.filter(
        payout_status='paid'
    ).aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    pending_amount = withdrawal_requests.filter(
        status__in=['new', 'processing']
    ).aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    available_amount = accruals_list.filter(
        payout_status='available'
    ).aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')
    payment_details = PartnerPaymentDetails.objects.filter(
        partner=request.user
    ).first()

    saved_payment_details = None

    if payment_details:
        saved_payment_details = {
            'card_number': decrypt_card_number(
                payment_details.encrypted_card_number
            ),
            'bank_name': payment_details.bank_name,
            'payment_system': payment_details.payment_system,
            'recipient_name': payment_details.recipient_name,
            'card_type': payment_details.card_type,
            'card_country': payment_details.card_country,
            'settlement_account': payment_details.settlement_account,
            'bik': payment_details.bik,
            'correspondent_account': payment_details.correspondent_account,
            'legal_bank_name': payment_details.legal_bank_name,
            'bank_inn': payment_details.bank_inn,
            'bank_kpp': payment_details.bank_kpp,
        }
    partner_registration = getattr(request.user, 'partner_registration', None)
    partner_type = partner_registration.partner_type if partner_registration else ''    
    return render(request, 'partner_portal/accruals.html', {
        'active_page': 'accruals',
        'accruals_list': accruals_list,
        'withdrawal_requests': withdrawal_requests,
        'total_accrued': total_accrued,
        'paid_amount': paid_amount,
        'pending_amount': pending_amount,
        'available_amount': available_amount,
        'saved_payment_details': saved_payment_details,
        'partner_type': partner_type,
        'new_accruals_list': new_accruals_list,
        'history_accruals_list': history_accruals_list,
    })

@login_required
@partner_required
def referral_program(request):
    links = PartnerLink.objects.filter(
        partner=request.user
    ).select_related('offer').order_by('-created_at')

    available_offers = Offer.objects.filter(
        partner_registrations__user=request.user,
        partner_registrations__is_approved=True
    ).distinct()

    for link in links:
        link.visits_count = link.visits.count()

        link.registrations_count = Lead.objects.filter(
            partner_link=link
        ).count()

        link.deals_count = Lead.objects.filter(
            partner_link=link,
            status='deal'
        ).count()

        link.deals_amount = (
            Lead.objects.filter(
                partner_link=link,
                status='deal'
            ).aggregate(total=Sum('partner_reward'))['total'] or 0
        )

    total_links = links.count()
    registrations_count = sum(link.registrations_count for link in links)
    total_deals = sum(link.deals_count for link in links)
    total_amount = sum(link.deals_amount for link in links)

    return render(request, 'partner_portal/referral_program.html', {
        'active_page': 'referral_program',
        'links': links,
        'available_offers': available_offers,
        'total_links': total_links,
        'total_deals': total_deals,
        'total_amount': total_amount,
        'registrations_count': registrations_count,
    })

@login_required
@partner_required
def create_partner_link(request):
    if request.method != 'POST':
        return redirect('partner_referral_program')

    offer_id = request.POST.get('offer')
    title = request.POST.get('title', '').strip() or 'Моя новая ссылка'

    offer = get_object_or_404(
        Offer,
        id=offer_id,
        partner_registrations__user=request.user,
        partner_registrations__is_approved=True
    )
    link_code = str(random.randint(1000000000, 9999999999))
    while PartnerLink.objects.filter(url__contains=f'link={link_code}').exists():
        link_code = str(random.randint(1000000000, 9999999999))
    partner_url = (
        f'{settings.SITE_URL}/users/client/register/'
        f'?referral={request.user.user_id}'
        f'&offer={offer.offer_id}'
        f'&link={link_code}'
    )
    PartnerLink.objects.create(
        partner=request.user,
        offer=offer,
        title=title,
        url=partner_url,
    )

    messages.success(request, 'Ссылка успешно добавлена.')
    return redirect('partner_referral_program')

@login_required
@partner_required
def request_withdrawal(request):
    if request.method != 'POST':
        return redirect('partner_accruals')

    comment = request.POST.get('comment', '').strip()

    available_accruals = Accrual.objects.filter(
        partner=request.user,
        payout_status='available'
    )
    card_number = request.POST.get('card_number', '').strip()
    encrypted_card_number = encrypt_card_number(card_number)
    card_number = request.POST.get('card_number', '')
    bank_name = request.POST.get('bank_name', '').strip()
    payment_system = request.POST.get('payment_system', '').strip()
    recipient_name = request.POST.get('recipient_name', '').strip()
    card_type = request.POST.get('card_type', '').strip()
    card_country = request.POST.get('card_country', '').strip()

    settlement_account = request.POST.get('settlement_account', '').strip()
    bik = request.POST.get('bik', '').strip()
    correspondent_account = request.POST.get('correspondent_account', '').strip()
    legal_bank_name = request.POST.get('legal_bank_name', '').strip()

    digits = ''.join(ch for ch in card_number if ch.isdigit())
    card_last4 = digits[-4:] if len(digits) >= 4 else ''
    card_mask = f'{digits[:4]} **** **** {card_last4}' if len(digits) >= 8 else ''
    bank_inn = request.POST.get('bank_inn', '').strip()
    bank_kpp = request.POST.get('bank_kpp', '').strip()
    PartnerPaymentDetails.objects.update_or_create(
    partner=request.user,
    defaults={
        'encrypted_card_number': encrypted_card_number,
        'card_mask': card_mask,
        'card_last4': card_last4,
        'bank_name': bank_name,
        'payment_system': payment_system,
        'recipient_name': recipient_name,
        'card_type': card_type,
        'card_country': card_country,
        'settlement_account': settlement_account,
        'bik': bik,
        'correspondent_account': correspondent_account,
        'legal_bank_name': legal_bank_name,
        'bank_inn': bank_inn,
        'bank_kpp': bank_kpp,
       }
    )
    total_amount = available_accruals.aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    if total_amount <= 0:
        messages.error(request, 'Нет доступных начислений для выплаты.')
        return redirect('partner_accruals')

    with transaction.atomic():
        withdrawal = WithdrawalRequest.objects.create(
            partner=request.user,
            admin=request.user.referred_by,
            amount=total_amount,
            status='new',
            recipient_name=recipient_name,
            card_mask=card_mask,
            card_last4=card_last4,
            bank_name=bank_name,
            payment_system=payment_system,
            card_type=card_type,
            card_country=card_country,
            comment=comment,
            encrypted_card_number=encrypted_card_number,
            settlement_account=settlement_account,
            bik=bik,
            correspondent_account=correspondent_account,
            legal_bank_name=legal_bank_name,
            bank_inn=bank_inn,
            bank_kpp=bank_kpp,
        )

        withdrawal.accruals.set(available_accruals)

        available_accruals.update(
            payout_status='requested'
        )

    messages.success(request, 'Заявка на выплату успешно отправлена.')
    return redirect('partner_accruals')

@login_required
@partner_required
def card_bin_lookup(request):
    bin_number = request.GET.get('bin', '')
    bin_number = ''.join(ch for ch in bin_number if ch.isdigit())[:8]

    if len(bin_number) < 6:
        return JsonResponse({
            'success': False,
            'message': 'Введите минимум 6 цифр карты.'
        })

    try:
        response = requests.get(
            'https://api.api-ninjas.com/v1/bin',
            params={'bin': bin_number},
            headers={'X-Api-Key': settings.API_NINJAS_KEY},
            timeout=5
        )
        response.raise_for_status()
        data = response.json()

        if not data:
            return JsonResponse({
                'success': False,
                'message': 'Данные по карте не найдены.'
            })

        card_data = data[0]

        return JsonResponse({
            'success': True,
            'bank': card_data.get('issuer', ''),
            'payment_system': card_data.get('brand', ''),
            'card_type': card_data.get('type', ''),
            'country': card_data.get('country', ''),
        })
    except requests.RequestException:
        return JsonResponse({
            'success': False,
            'message': 'Не удалось определить данные карты.'
        })

@login_required
@partner_required
def bank_bin_lookup(request):
    bin_value = request.GET.get('bin', '')
    bin_value = ''.join(ch for ch in bin_value if ch.isdigit())[:8]

    if len(bin_value) < 6:
        return JsonResponse({
            'success': False,
            'message': 'Введите минимум 6 цифр карты.'
        })

    bank_bin = BankBin.objects.filter(
        is_active=True,
        bin_prefix=bin_value[:6]
    ).first()

    if not bank_bin:
        return JsonResponse({
            'success': False,
            'message': 'Банк не найден.'
        })

    return JsonResponse({
        'success': True,
        'bank': bank_bin.bank_name,
        'logo': static(bank_bin.logo) if bank_bin.logo else '',
    })    
def encrypt_card_number(card_number):
    f = Fernet(settings.CARD_ENCRYPTION_KEY.encode())
    return f.encrypt(card_number.encode()).decode()


def decrypt_card_number(encrypted_card_number):
    f = Fernet(settings.CARD_ENCRYPTION_KEY.encode())
    return f.decrypt(encrypted_card_number.encode()).decode()

@login_required
@partner_required
def news_ajax(request):
    news = CompanyNews.objects.all()[:50]

    data = []

    for item in news:
        data.append({
            'id': item.id,
            'text': item.text,
            'published_at': timezone.localtime(item.published_at).strftime('%d.%m.%Y %H:%M') if item.published_at else '',
        })

    latest_id = news[0].id if news else None

    return JsonResponse({
        'latest_id': latest_id,
        'news': data,
    })

@login_required
@partner_required
def profile(request):
    partner_registration = getattr(request.user, 'partner_registration', None)

    payment_details = PartnerPaymentDetails.objects.filter(
        partner=request.user
    ).first()

    if request.method == 'POST':
        avatar = request.FILES.get('avatar')
        delete_avatar = request.POST.get('delete_avatar')

        if delete_avatar == '1':
            if request.user.avatar:
                request.user.avatar.delete(save=False)

            request.user.avatar = None
            request.user.save(update_fields=['avatar'])
            messages.success(request, 'Фото удалено')

        elif avatar:
            request.user.avatar = avatar
            request.user.save(update_fields=['avatar'])
            messages.success(request, 'Фото профиля сохранено')

        email = request.POST.get('email', '').strip()
        phone = request.POST.get('phone', '').strip()
        activity_type = request.POST.get('activity_type', '').strip()

        if email:
            request.user.email = email
            request.user.username = email

        request.user.phone = phone
        request.user.save(update_fields=['email', 'username', 'phone'])

        if partner_registration:

            partner_registration.email = email
            partner_registration.phone = phone
            partner_registration.activity_type = activity_type

            if partner_registration.partner_type == 'company':
                partner_registration.company_full_name = request.POST.get(
                    'company_full_name', ''
                ).strip()

                partner_registration.company_short_name = request.POST.get(
                    'company_short_name', ''
                ).strip()

                partner_registration.legal_address = request.POST.get(
                    'legal_address', ''
                ).strip()

                partner_registration.postal_address = request.POST.get(
                    'postal_address', ''
                ).strip()

                partner_registration.inn = request.POST.get(
                    'inn', ''
                ).strip()

                partner_registration.kpp = request.POST.get(
                    'kpp', ''
                ).strip()

                partner_registration.contact_person_name = request.POST.get(
                    'contact_person_name', ''
                ).strip()
            if partner_registration.partner_type == 'self_employed':
                partner_registration.full_name = request.POST.get('full_name', '').strip()
                partner_registration.activity_type = request.POST.get('activity_type', '').strip()
                partner_registration.company_name = request.POST.get('company_name', '').strip()
                partner_registration.inn = request.POST.get('inn', '').strip()
            partner_registration.save(update_fields=[
                'email',
                'phone',
                'activity_type',
                'full_name',
                'company_name',
                'company_full_name',
                'company_short_name',
                'legal_address',
                'postal_address',
                'inn',
                'kpp',
                'contact_person_name',
            ])

        card_number = request.POST.get('card_number', '').strip()
        bank_name = request.POST.get('bank_name', '').strip()
        payment_system = request.POST.get('payment_system', '').strip()
        bank_inn = request.POST.get('bank_inn', '').strip()
        bank_kpp = request.POST.get('bank_kpp', '').strip()
        legal_bank_name = request.POST.get('legal_bank_name', '').strip()
        bik = request.POST.get('bik', '').strip()
        settlement_account = request.POST.get('settlement_account', '').strip()
        correspondent_account = request.POST.get('correspondent_account', '').strip()
        if payment_details:
            if card_number:
                digits = ''.join(ch for ch in card_number if ch.isdigit())
                card_last4 = digits[-4:] if len(digits) >= 4 else ''
                card_mask = f'{digits[:4]} **** **** {card_last4}' if len(digits) >= 8 else ''
                payment_details.encrypted_card_number = encrypt_card_number(card_number)
                payment_details.card_mask = card_mask
                payment_details.card_last4 = card_last4
                payment_details.bank_name = bank_name
                payment_details.payment_system = payment_system
                payment_details.card_type = 'Дебетовая'
                payment_details.card_country = 'Россия'
            payment_details.legal_bank_name = legal_bank_name
            payment_details.bik = bik
            payment_details.settlement_account = settlement_account
            payment_details.correspondent_account = correspondent_account
            payment_details.bank_inn = bank_inn
            payment_details.bank_kpp = bank_kpp
            print('POST:', request.POST)
            print('bank_inn:', request.POST.get('bank_inn'))
            print('bank_kpp:', request.POST.get('bank_kpp'))
            payment_details.save(update_fields=[
                'encrypted_card_number',
                'card_mask',
                'card_last4',
                'bank_name',
                'payment_system',
                'card_type',
                'card_country',
                'legal_bank_name',
                'bik',
                'settlement_account',
                'correspondent_account',
                'bank_inn',
                'bank_kpp'
            ])

            messages.success(request, 'Реквизиты обновлены')

        return redirect('partner_profile')

    bank_logo = ''
    payment_logo = ''
    saved_card_number = ''

    if payment_details and payment_details.encrypted_card_number:
        saved_card_number = decrypt_card_number(payment_details.encrypted_card_number)

        bin_value = ''.join(ch for ch in saved_card_number if ch.isdigit())[:6]

        bank_bin = BankBin.objects.filter(
            is_active=True,
            bin_prefix=bin_value
        ).first()

        if bank_bin and bank_bin.logo:
            bank_logo = static(bank_bin.logo)

        payment_system = (payment_details.payment_system or '').lower()

        if payment_system == 'мир':
            payment_logo = static('users/images/payment/mir.svg')
        elif payment_system == 'visa':
            payment_logo = static('users/images/payment/visa.svg')
        elif payment_system == 'mastercard':
            payment_logo = static('users/images/payment/mastercard.svg')

    return render(request, 'partner_portal/profile.html', {
        'active_page': 'profile',
        'partner_registration': partner_registration,
        'payment_details': payment_details,
        'bank_logo': bank_logo,
        'payment_logo': payment_logo,
        'saved_card_number': saved_card_number,
    })

@login_required
@partner_required
def partner_leads(request):
    date_filter = request.GET.get('date_filter', 'all')
    selected_offer = request.GET.get('offer', '')
    selected_status = request.GET.get('status', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    search = request.GET.get('search', '')

    base_leads = Lead.objects.filter(
        partner=request.user
    ).select_related(
        'client',
        'offer',
        'partner_link',
        'tariff'
    ).order_by('-created_at')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    tabs_queryset = base_leads

    if selected_offer:
        tabs_queryset = tabs_queryset.filter(offer_id=selected_offer)

    if selected_status:
        tabs_queryset = tabs_queryset.filter(status=selected_status)

    if search:
        tabs_queryset = tabs_queryset.filter(
            Q(title__icontains=search) |
            Q(client__first_name__icontains=search) |
            Q(client__email__icontains=search) |
            Q(client__phone__icontains=search) |
            Q(offer__title__icontains=search)
        )

    count_all = tabs_queryset.count()
    count_today = tabs_queryset.filter(created_at__date=today).count()
    count_yesterday = tabs_queryset.filter(created_at__date=yesterday).count()
    count_week = tabs_queryset.filter(created_at__date__gte=week_start).count()
    count_month = tabs_queryset.filter(created_at__date__gte=month_start).count()

    if start_date or end_date:
        count_period_qs = tabs_queryset

        if start_date:
            count_period_qs = count_period_qs.filter(
                created_at__date__gte=start_date
            )

        if end_date:
            count_period_qs = count_period_qs.filter(
                created_at__date__lte=end_date
            )

        count_period = count_period_qs.count()
    else:
        count_period = 0

    leads_list = base_leads

    if date_filter == 'today':
        leads_list = leads_list.filter(created_at__date=today)
    elif date_filter == 'yesterday':
        leads_list = leads_list.filter(created_at__date=yesterday)
    elif date_filter == 'week':
        leads_list = leads_list.filter(created_at__date__gte=week_start)
    elif date_filter == 'month':
        leads_list = leads_list.filter(created_at__date__gte=month_start)
    elif date_filter == 'period':
        if start_date:
            leads_list = leads_list.filter(created_at__date__gte=start_date)
        if end_date:
            leads_list = leads_list.filter(created_at__date__lte=end_date)

    if selected_offer:
        leads_list = leads_list.filter(offer_id=selected_offer)

    if selected_status:
        leads_list = leads_list.filter(status=selected_status)

    if search:
        leads_list = leads_list.filter(
            Q(title__icontains=search) |
            Q(client__first_name__icontains=search) |
            Q(client__email__icontains=search) |
            Q(client__phone__icontains=search) |
            Q(offer__title__icontains=search)
        )

    total_leads = leads_list.count()
    new_leads = leads_list.filter(status='new').count()
    in_progress_leads = leads_list.filter(status='in_progress').count()
    cancelled_leads = leads_list.filter(status='cancelled').count()
    deal_leads = leads_list.filter(status='deal').count()

    conversion = round((deal_leads / total_leads) * 100) if total_leads else 0

    offers = Offer.objects.filter(
        partner_registrations__user=request.user,
        partner_registrations__is_approved=True
    ).annotate(
        leads_count=Count(
            'leads',
            filter=Q(leads__partner=request.user)
        )
    ).distinct()
    for lead in leads_list:
        lead.is_paid_out = lead.accruals.filter(
            payout_status='paid'
        ).exists()
    paginator = Paginator(leads_list, 20)
    page_number = request.GET.get('page')
    leads_page = paginator.get_page(page_number)

    query_params = request.GET.copy()
    query_params.pop('page', None)
    total_deal_amount = leads_list.aggregate(
        total=Sum('deal_amount')
    )['total'] or 0
    total_partner_reward = leads_list.aggregate(
        total=Sum('partner_reward')
    )['total'] or 0
    tariffs = Tariff.objects.all().order_by('price')
    partner_registrations = PartnerRegistration.objects.filter(
        user=request.user,
        is_approved=True
    ).select_related('offer', 'status')
    return render(request, 'partner_portal/leads.html', {
        'active_page': 'leads',

        'leads_list': leads_page,
        'pagination_query': query_params.urlencode(),

        'date_filter': date_filter,
        'selected_offer': selected_offer,
        'selected_status': selected_status,
        'start_date': start_date,
        'end_date': end_date,
        'search': search,

        'offers': offers,

        'count_all': count_all,
        'count_today': count_today,
        'count_yesterday': count_yesterday,
        'count_week': count_week,
        'count_month': count_month,
        'count_period': count_period,

        'count_status_all': total_leads,
        'count_status_new': new_leads,
        'count_status_in_progress': in_progress_leads,
        'count_status_cancelled': cancelled_leads,
        'count_status_deal': deal_leads,

        'total_leads': total_leads,
        'new_leads': new_leads,
        'in_progress_leads': in_progress_leads,
        'cancelled_leads': cancelled_leads,
        'deal_leads': deal_leads,
        'conversion': conversion,
        'total_deal_amount': total_deal_amount,
        'total_partner_reward': total_partner_reward,
        'tariffs': tariffs,
        'partner_registrations': partner_registrations,
    })
@login_required
@partner_required
def partner_create_lead(request):
    if request.method != 'POST':
        return redirect('partner_leads')

    title = request.POST.get('title', 'Новый лид').strip()
    offer_id = request.POST.get('offer')
    tariff_id = request.POST.get('tariff')

    client_name = request.POST.get('client_name', '').strip()
    client_email = request.POST.get('client_email', '').strip()
    client_phone = request.POST.get('client_phone', '').strip()
    client_company = request.POST.get('client_company', '').strip()

    offer = Offer.objects.filter(id=offer_id).first()
    tariff = Tariff.objects.filter(id=tariff_id).first()

    if not offer:
        messages.error(request, 'Выберите оффер.')
        return redirect('partner_leads')

    if not tariff:
        messages.error(request, 'Выберите тариф.')
        return redirect('partner_leads')

    if not client_name or not client_email or not client_phone:
        messages.error(request, 'Заполните данные клиента.')
        return redirect('partner_leads')

    alphabet = string.ascii_letters + string.digits
    temporary_password = ''.join(secrets.choice(alphabet) for _ in range(10))

    UserModel = get_user_model()

    existing_client = UserModel.objects.filter(
        email=client_email
    ).first()

    if existing_client:
        messages.error(
            request,
            'Клиент с таким email уже существует. Введите другой email.'
        )
        return redirect('partner_leads')

    client, created = UserModel.objects.get_or_create(
        email=client_email,
        defaults={
            'username': client_email,
            'first_name': client_name,
            'phone': client_phone,
            'company': client_company,
        }
    )

    if created:
        client.set_password(temporary_password)
        client.referred_by = request.user
        client.save()
    else:
        temporary_password = 'Клиент уже был зарегистрирован ранее'

    Lead.objects.create(
        title=title or 'Новый лид',
        deal_amount=0,
        partner_reward=0,
        partner=request.user,
        client=client,
        status='new',
        offer=offer,
        tariff=tariff,
        admin=request.user.referred_by,
        partner_link=None,
    )

    messages.success(
        request,
        f'Лид добавлен. Клиент зарегистрирован. Временный пароль: {temporary_password}'
    )

    return redirect('partner_leads')
@login_required
@partner_required
def partner_update_lead(request, lead_id):
    if request.method != 'POST':
        return redirect('partner_leads')

    lead = get_object_or_404(
        Lead,
        id=lead_id,
        partner=request.user
    )

    if lead.accruals.filter(payout_status='paid').exists():
        messages.error(request, 'Лид нельзя редактировать: по нему уже произведена выплата.')
        return redirect('partner_leads')

    title = request.POST.get('title', '').strip()
    created_at = request.POST.get('created_at', '').strip()
    offer_id = request.POST.get('offer')
    status = request.POST.get('status')
    tariff_id = request.POST.get('tariff')

    client_name = request.POST.get('client_name', '').strip()
    client_email = request.POST.get('client_email', '').strip()
    client_phone = request.POST.get('client_phone', '').strip()
    client_company = request.POST.get('client_company', '').strip()

    offer = Offer.objects.filter(id=offer_id).first()
    tariff = Tariff.objects.filter(id=tariff_id).first()

    if not title or not offer or not tariff or not client_name or not client_email or not client_phone:
        messages.error(request, 'Заполните все обязательные поля.')
        return redirect('partner_leads')

    if status == 'deal':
        messages.error(request, 'Партнер не может установить статус "Сделка".')
        return redirect('partner_leads')

    lead.title = title
    lead.offer = offer
    lead.status = status
    lead.tariff = tariff

    if lead.client:
        lead.client.first_name = client_name
        lead.client.email = client_email
        lead.client.username = client_email
        lead.client.phone = client_phone
        lead.client.company = client_company
        lead.client.save(update_fields=[
            'first_name',
            'email',
            'username',
            'phone',
            'company',
        ])

    lead.save(update_fields=[
        'title',
        'offer',
        'status',
        'tariff',
    ])

    messages.success(request, 'Лид успешно обновлен.')
    return redirect('partner_leads')

@login_required
@partner_required
def partner_reports(request):
    date_filter = request.GET.get('date_filter', 'all')
    report_type = request.GET.get('report_type', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    generate_report = request.GET.get('generate_report')

    leads_report = None
    report_generated = False

    base_leads = Lead.objects.filter(
        partner=request.user
    ).select_related(
        'client',
        'offer',
        'tariff'
    ).order_by('-created_at')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    if generate_report:
        report_generated = True
        leads_report = base_leads

        if date_filter == 'today':
            leads_report = leads_report.filter(created_at__date=today)
        elif date_filter == 'yesterday':
            leads_report = leads_report.filter(created_at__date=yesterday)
        elif date_filter == 'week':
            leads_report = leads_report.filter(created_at__date__gte=week_start)
        elif date_filter == 'month':
            leads_report = leads_report.filter(created_at__date__gte=month_start)
        elif date_filter == 'period':
            if start_date:
                leads_report = leads_report.filter(created_at__date__gte=start_date)
            if end_date:
                leads_report = leads_report.filter(created_at__date__lte=end_date)

        if report_type == 'deals':
            leads_report = leads_report.filter(status='deal')
        elif report_type == 'no_deals':
            leads_report = leads_report.exclude(status='deal')

        total_leads = leads_report.count()
        deal_leads = leads_report.filter(status='deal').count()

        total_amount = leads_report.aggregate(
            total=Sum('deal_amount')
        )['total'] or 0

        total_reward = leads_report.aggregate(
            total=Sum('partner_reward')
        )['total'] or 0

        paginator = Paginator(leads_report, 20)
        page_number = request.GET.get('page')
        leads_report = paginator.get_page(page_number)

    else:
        total_leads = 0
        deal_leads = 0
        total_amount = 0
        total_reward = 0

    period_text = ''

    if date_filter == 'today':
        period_text = f'за {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'yesterday':
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'
    elif date_filter == 'week':
        period_text = f'за период с {week_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'month':
        period_text = f'за период с {month_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'period':
        if start_date and end_date:
            period_text = f'за период с {start_date} по {end_date}'
        elif start_date:
            period_text = f'с {start_date}'
        elif end_date:
            period_text = f'по {end_date}'
    partner_reg = PartnerRegistration.objects.filter(
        user=request.user,
        is_approved=True
    ).first()

    if partner_reg and partner_reg.partner_type == 'company':
        partner_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or request.user.email
        )
    else:
        partner_name = (
            f'{request.user.last_name} '
            f'{request.user.first_name} '
            f'{getattr(request.user, "middle_name", "") or ""}'
        ).strip() or request.user.email
    return render(request, 'partner_portal/reports.html', {
        'active_page': 'reports',

        'leads_report': leads_report,
        'report_generated': report_generated,

        'date_filter': date_filter,
        'report_type': report_type,
        'start_date': start_date,
        'end_date': end_date,
        'period_text': period_text,

        'total_leads': total_leads,
        'deal_leads': deal_leads,
        'total_amount': total_amount,
        'total_reward': total_reward,
        'partner_name': partner_name,
    })
@login_required
@partner_required
def export_partner_reports_word(request):
    date_filter = request.GET.get('date_filter', 'all')
    report_type = request.GET.get('report_type', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    leads = Lead.objects.filter(
        partner=request.user
    ).select_related(
        'client',
        'offer',
        'tariff'
    ).order_by('-created_at')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)
    formatted_start = ''
    formatted_end = ''

    if start_date:
        formatted_start = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d.%m.%Y')

    if end_date:
        formatted_end = datetime.strptime(end_date, '%Y-%m-%d').strftime('%d.%m.%Y')
    period_text = ''

    if date_filter == 'today':
        leads = leads.filter(created_at__date=today)
        period_text = f'за {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'yesterday':
        leads = leads.filter(created_at__date=yesterday)
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'

    elif date_filter == 'week':
        leads = leads.filter(created_at__date__gte=week_start)
        period_text = f'за период с {week_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'month':
        leads = leads.filter(created_at__date__gte=month_start)
        period_text = f'за период с {month_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'period':
        if start_date:
            leads = leads.filter(created_at__date__gte=start_date)
        if end_date:
            leads = leads.filter(created_at__date__lte=end_date)

        if start_date and end_date:
            period_text = f'за период с {formatted_start} по {formatted_end}'
        elif start_date:
            period_text = f'с {start_date}'
        elif end_date:
            period_text = f'по {end_date}'

    if report_type == 'deals':
        leads = leads.filter(status='deal')
    elif report_type == 'no_deals':
        leads = leads.exclude(status='deal')

    total_leads = leads.count()
    deal_leads = leads.filter(status='deal').count()

    total_amount = leads.aggregate(
        total=Sum('deal_amount')
    )['total'] or 0

    total_reward = leads.aggregate(
        total=Sum('partner_reward')
    )['total'] or 0

    document = Document()
    style = document.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width, section.page_height = section.page_height, section.page_width

    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

    partner_reg = PartnerRegistration.objects.filter(
        user=request.user,
        is_approved=True
    ).first()

    if partner_reg and partner_reg.partner_type == 'company':
        partner_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or request.user.email
        )
    else:
        partner_name = (
            f'{request.user.last_name} '
            f'{request.user.first_name} '
            f'{getattr(request.user, "middle_name", "") or ""}'
        ).strip() or request.user.email

    title_text = f'Отчёт по лидам партнёра: {partner_name}'

    if period_text:
        title_text += f'\n{period_text}'

    title = document.add_heading(title_text, level=1)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for run in title.runs:
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph('')

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.autofit = False

    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')

    for margin in ['top', 'left', 'bottom', 'right']:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), "30")
        node.set(qn('w:type'), 'dxa')
        tblCellMar.append(node)

    tblPr.append(tblCellMar)
    table.style = 'Table Grid'
    column_widths = [
        Cm(2.1),  # Дата
        Cm(1.8),  # Лид
        Cm(3.7),  # Клиент
        Cm(4.2),  # Оффер
        Cm(2.0),  # Статус
        Cm(2.1),  # Сумма сделки
        Cm(2.1),  # Вознаграждение
    ]
    table.autofit = False
    headers = [
        'Дата',
        'Лид',
        'Клиент',
        'Оффер',
        'Статус',
        'Сумма сделки',
        'Вознаграждение',
    ]
    for row in table.rows:
        for idx, width in enumerate(column_widths):
            row.cells[idx].width = width
    header_cells = table.rows[0].cells

    for index, header in enumerate(headers):
        cell = header_cells[index]

        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)

        run.bold = True
        run.font.size = Pt(9)

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for lead in leads:
        row_cells = table.add_row().cells

        row_cells[0].text = lead.created_at.strftime('%d.%m.%Y %H:%M')
        row_cells[1].text = lead.title or '—'

        if lead.client:
            row_cells[2].text = f'{lead.client.first_name or "—"}\n{lead.client.email or ""}'
        else:
            row_cells[2].text = '—'

        row_cells[3].text = lead.offer.title if lead.offer else '—'
        row_cells[4].text = lead.get_status_display()
        row_cells[5].text = f'{int(lead.deal_amount or 0):,} ₽'.replace(',', ' ')
        row_cells[6].text = f'{int(lead.partner_reward or 0):,} ₽'.replace(',', ' ')

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7)
        for idx, cell in enumerate(row_cells):
            for paragraph in cell.paragraphs:

                if idx in [5, 6]:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                else:
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                for run in paragraph.runs:
                    run.font.size = Pt(7)
    document.add_paragraph()

    summary_table = document.add_table(rows=1, cols=4)
    summary_table.autofit = False

    summary_table.columns[0].width = Inches(1.7)
    summary_table.columns[1].width = Inches(1.7)
    summary_table.columns[2].width = Inches(2.0)
    summary_table.columns[3].width = Inches(2.0)

    summary_data = [
        f'Лидов: {total_leads}',
        f'Сделок: {deal_leads}',
        f'Сумма сделок: {total_amount:,} ₽'.replace(',', ' '),
        f'Вознаграждение: {total_reward:,} ₽'.replace(',', ' ')
    ]

    for i, text in enumerate(summary_data):
        cell = summary_table.rows[0].cells[i]

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(10)

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        borders = OxmlElement('w:tcBorders')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            borders.append(border)

    tcPr.append(borders)
    document.add_paragraph('')
    section = document.sections[0]
    section.page_height
    last_paragraph = document.paragraphs[-1]
    last_paragraph.paragraph_format.space_after = Pt(120)
    signatures = document.add_table(rows=1, cols=3)
    signatures.autofit = False

    signatures.columns[0].width = Inches(2.8)
    signatures.columns[1].width = Inches(2.8)
    signatures.columns[2].width = Inches(1.8)

    # Заказчик
    cell = signatures.rows[0].cells[0]
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    run = p.add_run('Заказчик ____________________')
    run.bold = True
    run.font.size = Pt(9)

    # Исполнитель
    cell = signatures.rows[0].cells[1]
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = p.add_run('Исполнитель ____________________')
    run.bold = True
    run.font.size = Pt(9)

    # Дата
    cell = signatures.rows[0].cells[2]
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = p.add_run(f'Дата: {today.strftime("%d.%m.%Y")}')
    run.bold = True
    run.font.size = Pt(9)
    # убрать границы
    for row in signatures.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            borders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                borders.append(border)

            tcPr.append(borders)
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename="partner_report.docx"'

    document.save(response)

    return response
@login_required
@partner_required
def export_partner_reports_pdf(request):
    date_filter = request.GET.get('date_filter', 'all')
    report_type = request.GET.get('report_type', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    leads = Lead.objects.filter(
        partner=request.user
    ).select_related(
        'client',
        'offer',
        'tariff'
    ).order_by('-created_at')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    formatted_start = ''
    formatted_end = ''

    if start_date:
        formatted_start = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d.%m.%Y')

    if end_date:
        formatted_end = datetime.strptime(end_date, '%Y-%m-%d').strftime('%d.%m.%Y')

    period_text = ''

    if date_filter == 'today':
        leads = leads.filter(created_at__date=today)
        period_text = f'за {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'yesterday':
        leads = leads.filter(created_at__date=yesterday)
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'
    elif date_filter == 'week':
        leads = leads.filter(created_at__date__gte=week_start)
        period_text = f'за период с {week_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'month':
        leads = leads.filter(created_at__date__gte=month_start)
        period_text = f'за период с {month_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'period':
        if start_date:
            leads = leads.filter(created_at__date__gte=start_date)
        if end_date:
            leads = leads.filter(created_at__date__lte=end_date)

        if start_date and end_date:
            period_text = f'за период с {formatted_start} по {formatted_end}'
        elif start_date:
            period_text = f'с {formatted_start}'
        elif end_date:
            period_text = f'по {formatted_end}'

    if report_type == 'deals':
        leads = leads.filter(status='deal')
    elif report_type == 'no_deals':
        leads = leads.exclude(status='deal')

    total_leads = leads.count()
    deal_leads = leads.filter(status='deal').count()

    total_amount = leads.aggregate(total=Sum('deal_amount'))['total'] or 0
    total_reward = leads.aggregate(total=Sum('partner_reward'))['total'] or 0

    partner_reg = PartnerRegistration.objects.filter(
        user=request.user,
        is_approved=True
    ).first()

    if partner_reg and partner_reg.partner_type == 'company':
        partner_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or request.user.email
        )
    else:
        partner_name = (
            f'{request.user.last_name} '
            f'{request.user.first_name} '
            f'{getattr(request.user, "middle_name", "") or ""}'
        ).strip() or request.user.email

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="partner_report.pdf"'

    pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))

    doc = SimpleDocTemplate(
        response,
        pagesize=landscape(A4),
        rightMargin=8 * mm,
        leftMargin=8 * mm,
        topMargin=8 * mm,
        bottomMargin=8 * mm
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'ReportTitle',
        parent=styles['Title'],
        fontName='Arial-Bold',
        fontSize=12,
        leading=14,
        alignment=1,
        textColor=colors.HexColor('#000'),
    )

    normal_style = ParagraphStyle(
        'NormalRu',
        parent=styles['Normal'],
        fontName='Arial',
        fontSize=7,
        leading=8,
    )

    story = []

    title = f'Отчёт по лидам партнёра: {partner_name}'
    if period_text:
        title += f'<br/>{period_text}'

    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 6))

    data = [[
        'Дата',
        'Лид',
        'Клиент',
        'Оффер',
        'Статус',
        'Сумма сделки',
        'Вознаграждение',
    ]]

    for lead in leads:
        client_text = '—'
        if lead.client:
            client_text = f'{lead.client.first_name or "—"}<br/>{lead.client.email or ""}'

        data.append([
            lead.created_at.strftime('%d.%m.%Y %H:%M'),
            lead.title or '—',
            Paragraph(client_text, normal_style),
            Paragraph(lead.offer.title if lead.offer else '—', normal_style),
            lead.get_status_display(),
            f'{int(lead.deal_amount or 0):,} ₽'.replace(',', ' '),
            f'{int(lead.partner_reward or 0):,} ₽'.replace(',', ' '),
        ])

    table = Table(
        data,
        colWidths=[
            30 * mm,  # Дата
            25 * mm,  # Лид
            50 * mm,  # Клиент
            60 * mm,  # Оффер
            30 * mm,  # Статус
            35 * mm,  # Сумма сделки
            40 * mm,  # Вознаграждение
        ],
        repeatRows=1
    )

    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, 0), 'Arial-Bold'),
        ('FONTNAME', (0, 1), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, -1), 6),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('ALIGN', (5, 1), (6, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.black),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
    ]))

    story.append(table)
    story.append(Spacer(1, 8))

    summary_data = [[
        f'Лидов: {total_leads}',
        f'Сделок: {deal_leads}',
        f'Сумма сделок: {int(total_amount):,} ₽'.replace(',', ' '),
        f'Вознаграждение: {int(total_reward):,} ₽'.replace(',', ' '),
    ]]

    summary_table = Table(
        summary_data,
        colWidths=[
            65 * mm,
            65 * mm,
            70 * mm,
            70 * mm,
        ]
    )

    summary_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Arial-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOX', (0, 0), (-1, -1), 0, colors.white),
        ('INNERGRID', (0, 0), (-1, -1), 0, colors.white),
    ]))

    story.append(summary_table)
    story.append(Spacer(1, 8))

    summary_data = [[
        f'Лидов: {total_leads}',
        f'Сделок: {deal_leads}',
        f'Сумма сделок: {int(total_amount):,} ₽'.replace(',', ' '),
        f'Вознаграждение: {int(total_reward):,} ₽'.replace(',', ' '),
    ]]
 
    def draw_footer(canvas, doc):
        width, height = landscape(A4)

        canvas.saveState()
        canvas.setFont('Arial-Bold', 8)

        # Левая часть
        canvas.drawString(
            10 * mm,
            12 * mm,
            'Заказчик ____________________'
        )

        # Центр
        canvas.drawCentredString(
            width / 2,
            12 * mm,
            'Исполнитель ____________________'
        )

        # Правая часть
        canvas.drawRightString(
            width - 10 * mm,
            12 * mm,
            f'Дата: {today.strftime("%d.%m.%Y")}'
        )

        # Номер страницы снизу по центру
        canvas.drawCentredString(
            width / 2,
            7 * mm,
            f'Страница {doc.page}'
        )

        canvas.restoreState()
    doc.build(
        story,
        onFirstPage=draw_footer,
        onLaterPages=draw_footer
    )

    return response
@login_required
@partner_required
def export_partner_reports_csv(request):
    date_filter = request.GET.get('date_filter', 'all')
    report_type = request.GET.get('report_type', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    leads = Lead.objects.filter(
        partner=request.user
    ).select_related(
        'client',
        'offer',
        'tariff'
    ).order_by('-created_at')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    period_text = ''

    if date_filter == 'today':
        leads = leads.filter(created_at__date=today)
        period_text = f'за {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'yesterday':
        leads = leads.filter(created_at__date=yesterday)
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'

    elif date_filter == 'week':
        leads = leads.filter(created_at__date__gte=week_start)
        period_text = (
            f'за период с '
            f'{week_start.strftime("%d.%m.%Y")} '
            f'по {today.strftime("%d.%m.%Y")}'
        )

    elif date_filter == 'month':
        leads = leads.filter(created_at__date__gte=month_start)
        period_text = (
            f'за период с '
            f'{month_start.strftime("%d.%m.%Y")} '
            f'по {today.strftime("%d.%m.%Y")}'
        )

    elif date_filter == 'period':
        if start_date:
            leads = leads.filter(created_at__date__gte=start_date)

        if end_date:
            leads = leads.filter(created_at__date__lte=end_date)

        if start_date and end_date:
            formatted_start = datetime.strptime(
                start_date,
                '%Y-%m-%d'
            ).strftime('%d.%m.%Y')

            formatted_end = datetime.strptime(
                end_date,
                '%Y-%m-%d'
            ).strftime('%d.%m.%Y')

            period_text = (
                f'за период с {formatted_start} '
                f'по {formatted_end}'
            )

    if report_type == 'deals':
        leads = leads.filter(status='deal')

    elif report_type == 'no_deals':
        leads = leads.exclude(status='deal')

    total_leads = leads.count()

    deal_leads = leads.filter(
        status='deal'
    ).count()

    total_amount = leads.aggregate(
        total=Sum('deal_amount')
    )['total'] or 0

    total_reward = leads.aggregate(
        total=Sum('partner_reward')
    )['total'] or 0

    response = HttpResponse(
        content_type='text/csv; charset=utf-8'
    )

    response['Content-Disposition'] = (
        'attachment; filename="partner_report.csv"'
    )

    # BOM для Excel
    response.write('\ufeff')

    writer = csv.writer(
        response,
        delimiter=';'
    )

    partner_reg = PartnerRegistration.objects.filter(
        user=request.user,
        is_approved=True
    ).first()

    if partner_reg and partner_reg.partner_type == 'company':
        partner_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or request.user.email
        )
    else:
        partner_name = (
            f'{request.user.last_name} '
            f'{request.user.first_name} '
            f'{getattr(request.user, "middle_name", "") or ""}'
        ).strip() or request.user.email

    writer.writerow([
        f'Отчёт по лидам партнёра: {partner_name}'
    ])

    if period_text:
        writer.writerow([period_text])

    writer.writerow([])

    writer.writerow([
        'Дата',
        'Лид',
        'Клиент',
        'Email клиента',
        'Оффер',
        'Статус',
        'Сумма сделки',
        'Вознаграждение',
    ])

    for lead in leads:
        writer.writerow([
            lead.created_at.strftime('%d.%m.%Y %H:%M'),
            lead.title or '—',
            (
                lead.client.first_name
                if lead.client else '—'
            ),
            (
                lead.client.email
                if lead.client else '—'
            ),
            (
                lead.offer.title
                if lead.offer else '—'
            ),
            lead.get_status_display(),
            f'{int(lead.deal_amount or 0):,} ₽'.replace(',', ' '),
            f'{int(lead.partner_reward or 0):,} ₽'.replace(',', ' '),
        ])

    writer.writerow([])
    writer.writerow(['ИТОГИ'])

    writer.writerow([
        f'Лидов: {total_leads}',
        f'Сделок: {deal_leads}',
        f'Сумма сделок: {int(total_amount):,} ₽'.replace(',', ' '),
        f'Вознаграждение: {int(total_reward):,} ₽'.replace(',', ' '),
    ])

    return response
@login_required
@partner_required
def partner_education(request):

    videos = EducationVideo.objects.order_by('-created_at')
    documents = EducationDocument.objects.order_by('-created_at')

    context = {
        'videos': videos,
        'documents': documents,
        'active_page': 'education',
    }

    return render(
        request,
        'partner_portal/education.html',
        context
    )
@login_required
@partner_required
def partner_materials(request):

    partner_offers = Offer.objects.filter(
        partner_registrations__user=request.user,
        partner_registrations__is_approved=True
    ).distinct()

    videos = MarketingMaterial.objects.filter(
        material_type='video',
        offer__in=partner_offers
    ).order_by('-created_at')

    images = MarketingMaterial.objects.filter(
        material_type='image',
        offer__in=partner_offers
    ).order_by('-created_at')

    texts = MarketingMaterial.objects.filter(
        material_type='text',
        offer__in=partner_offers
    ).order_by('-created_at')

    context = {
        'videos': videos,
        'images': images,
        'texts': texts,
        'active_page': 'materials',
    }

    return render(
        request,
        'partner_portal/materials.html',
        context
    )
@login_required
@partner_required
def partner_chat(request):
    admin_user = User.objects.filter(
        is_staff=True
    ).first()
    topics = (
        ChatTopic.objects.filter(partner=request.user)
        .select_related('admin')
        .prefetch_related('messages')
        .annotate(
            unread_count=Count(
                'messages',
                filter=Q(messages__is_read=False) & ~Q(messages__sender=request.user)
            )
        )
        .order_by('-updated_at')
    )

    return render(request, 'partner_portal/partner_chat.html', {
        'active_page': 'partner_chat',
        'topics': topics,
        'admin_user': admin_user,
    })
@login_required
@partner_required
def partner_chat_topic(request, topic_id):
    admin_user = User.objects.filter(
    is_staff=True
    ).first()
    topics = (
        ChatTopic.objects.filter(partner=request.user)
        .select_related('admin')
        .prefetch_related('messages')
        .annotate(
            unread_count=Count(
                'messages',
                filter=Q(messages__is_read=False) & ~Q(messages__sender=request.user)
            )
        )
        .order_by('-updated_at')
    )
    current_topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )
    current_topic.messages.filter(
        is_read=False
    ).exclude(
        sender=request.user
    ).update(
        is_read=True
    )
    messages = (
        current_topic.messages
        .select_related('sender')
        .order_by('created_at')
    )

    return render(request, 'partner_portal/partner_chat.html', {
        'active_page': 'partner_chat',
        'topics': topics,
        'current_topic': current_topic,
        'messages': messages,
        'admin_user': admin_user,
    })

@login_required
@partner_required
def partner_chat_create_topic(request):
    if request.method == 'POST':
        title = request.POST.get('title', '').strip()

        if title:
            admin = request.user.referred_by

            topic = ChatTopic.objects.create(
                title=title,
                partner=request.user,
                admin=admin,
                created_by=request.user,
                status='open',
                priority='normal',
                category='other'
            )

            return redirect('partner_chat_topic', topic_id=topic.id)

    return redirect('partner_chat')

@login_required
@partner_required
def partner_chat_send_message(request, topic_id):

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    if topic.status == 'closed':
        return JsonResponse({
            'success': False,
            'error': 'Тема закрыта'
        })

    if request.method == 'POST':

        text = request.POST.get('text', '').strip()

        if text:

            message = ChatMessage.objects.create(
                topic=topic,
                sender=request.user,
                text=text
            )
            request.user.typing_updated_at = None
            request.user.save(update_fields=['typing_updated_at'])
            topic.save()

            if request.headers.get('x-requested-with') == 'XMLHttpRequest':

                return JsonResponse({
                    'success': True,
                    'id': message.id,
                    'text': message.text,
                    'time': message.created_at.strftime('%H:%M'),
                    'is_own': True,
                })

    return redirect('partner_chat_topic', topic_id=topic.id)

@login_required
@partner_required
def partner_chat_messages(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    after_id = request.GET.get('after_id', 0)

    try:
        after_id = int(after_id)
    except (TypeError, ValueError):
        after_id = 0

    messages = (
        topic.messages
        .select_related('sender')
        .filter(id__gt=after_id)
        .order_by('created_at')
    )

    data = []

    for message in messages:
        data.append({
            'id': message.id,
            'text': message.text,
            'time': message.created_at.strftime('%H:%M'),
            'is_read': message.is_read,
            'is_own': message.sender_id == request.user.id,
            'message_type': message.message_type,
            'audio_url': message.audio_file.url if message.audio_file else None,
            'image_url': message.image_file.url if message.image_file else None,
            'video_note_url': message.video_note.url if message.video_note else None,
            'file_url': message.file.url if message.file else None,
            'file_name': message.file.name.split('/')[-1] if message.file else None,
            'avatar_url': message.sender.avatar.url if message.sender.avatar else '/static/img/default-avatar.png',
        })
    typing = topic.admin.is_typing if topic.admin else False

    return JsonResponse({
        'messages': data,
        'typing': typing,
    })

@login_required
@partner_required
def partner_chat_typing(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    if request.method == 'POST':
        request.user.typing_updated_at = timezone.now()
        request.user.save(update_fields=['typing_updated_at'])

    return JsonResponse({'success': True})
@login_required
@partner_required
def partner_chat_latest_unread_topic(request):
    current_topic_id = request.GET.get('current_topic_id')

    current_topic = ChatTopic.objects.filter(
        id=current_topic_id,
        partner=request.user
    ).first()

    topics = (
        ChatTopic.objects.filter(partner=request.user)
        .prefetch_related('messages')
        .order_by('-updated_at')
    )

    for topic in topics:
        if current_topic and topic.updated_at <= current_topic.updated_at:
            continue

        last_message = topic.messages.order_by('-created_at').first()

        if not last_message:
            continue

        if last_message.sender_id == request.user.id:
            continue

        return JsonResponse({
            'has_topic': True,
            'topic_id': topic.id,
            'url': reverse('partner_chat_topic', args=[topic.id]),
        })

    return JsonResponse({'has_topic': False})
@login_required
@partner_required
def partner_chat_send_voice(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    if request.method == 'POST':
        audio = request.FILES.get('audio')

        if audio:
            message = ChatMessage.objects.create(
                topic=topic,
                sender=request.user,
                message_type='voice',
                audio_file=audio,
                text=''
            )

            topic.updated_at = timezone.now()
            topic.status = 'open'
            topic.save(update_fields=['updated_at', 'status'])

            return JsonResponse({
                'success': True,
                'id': message.id,
                'message_type': message.message_type,
                'audio_url': message.audio_file.url,
                'time': message.created_at.strftime('%H:%M'),
                'is_own': True,
            })

    return JsonResponse({
        'success': False
    })
@login_required
@partner_required
def partner_admin_status(request, topic_id):

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    admin = topic.admin

    return JsonResponse({
        'is_online': admin.is_online,
        'last_seen_text': admin.last_seen_text,
        'is_typing': admin.is_typing,
    })
@login_required
@partner_required
def chat_read_statuses(request, topic_id):

    topic = get_object_or_404(ChatTopic, id=topic_id)

    messages = topic.messages.filter(
        sender=request.user
    )

    data = []

    for message in messages:
        data.append({
            'id': message.id,
            'is_read': message.is_read,
        })

    return JsonResponse({
        'messages': data
    })
@login_required
@partner_required
def partner_chat_mark_read(request, topic_id):

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    if request.method == 'POST':

        topic.messages.filter(
            is_read=False
        ).exclude(
            sender=request.user
        ).update(
            is_read=True
        )

    return JsonResponse({
        'success': True
    })
@login_required
@partner_required
def partner_chat_send_video_note(request, topic_id):
    if request.method != 'POST':
        return JsonResponse({'success': False})

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    video = request.FILES.get('video')

    if not video:
        return JsonResponse({'success': False})

    message = ChatMessage.objects.create(
        topic=topic,
        sender=request.user,
        message_type='video_note',
        video_note=video
    )

    return JsonResponse({
        'success': True,
        'id': message.id,
        'message_type': message.message_type,
        'video_note_url': message.video_note.url,
        'time': message.created_at.strftime('%H:%M'),
        'is_own': True,
        'is_read': False,
    })
@login_required
@partner_required
def partner_chat_toggle_topic_status(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        partner=request.user
    )

    if request.method == 'POST':
        if topic.status == 'closed':
            topic.status = 'open'
        else:
            topic.status = 'closed'

        topic.updated_at = timezone.now()
        topic.save(update_fields=['status', 'updated_at'])

    return redirect('partner_chat_topic', topic_id=topic.id)