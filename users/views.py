from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib.auth.views import LoginView, LogoutView
from django.shortcuts import get_object_or_404, redirect, render
from datetime import timedelta
from django.utils import timezone
from django.db import models
from django.urls import reverse, reverse_lazy
from urllib.parse import urlencode 
from offers.models import Offer, Lead, OfferPromoMaterial, PartnerRegistration, PartnerStatus, OfferVisit, PartnerLink, Lead, Tariff
from offers.forms import OfferCreateForm, OfferUpdateForm, PartnerRegistrationForm, PartnerCreateByUserForm
from offers.models import Offer, OfferPromoMaterial, PartnerRegistration, PartnerStatus
from .forms import AdminLoginForm, AdminUserRegisterForm, ProfileUpdateForm, ClientRegistrationForm
from django.utils.crypto import get_random_string
from django.contrib.auth import get_user_model, update_session_auth_hash
from django.db.models import Q, Exists, OuterRef,Count, Sum, DecimalField, Value, F, Subquery, IntegerField
from django.db.models.functions import Coalesce
from django.contrib import messages
from django.conf import settings
from users.models import User, Accrual, Subscription, WithdrawalRequest, MarketingMaterial, EducationVideo, PartnerPaymentDetails, EducationDocument, ChatTopic, ChatMessage
from decimal import Decimal
from datetime import datetime
from django.db import transaction
from django.core.paginator import Paginator
from django.http import JsonResponse, HttpResponse
from cryptography.fernet import Fernet
import csv
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.shared import Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.lib.enums import TA_CENTER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from functools import wraps
from moviepy import VideoFileClip
import os
from django.db.models.functions import TruncMonth
from ai_service.whisper import transcribe_audio
from asgiref.sync import async_to_sync
from channels.layers import get_channel_layer


def admin_required(view_func):
    @wraps(view_func)
    def wrapper(request, *args, **kwargs):
        if not request.user.is_authenticated:
            return redirect('login')

        if request.user.role != 'admin' and not request.user.is_staff:
            return redirect('partner_dashboard')

        return view_func(request, *args, **kwargs)

    return wrapper

def save_offer_files(offer, files):
    existing_count = offer.promo_files.count()
    remaining_slots = max(0, 10 - existing_count)

    for file_obj in files[:remaining_slots]:

        promo = OfferPromoMaterial.objects.create(
            offer=offer,
            file=file_obj,
        )

        extension = file_obj.name.split('.')[-1].lower()

        if extension in ['mp4', 'mov', 'avi', 'webm']:
            material_type = 'video'

        elif extension in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
            material_type = 'image'

        else:
            material_type = 'text'

        MarketingMaterial.objects.create(
            title=file_obj.name,
            description=f'Промоматериал оффера: {offer.title}',
            material_type=material_type,
            file=promo.file,
            offer=offer,
        )
        
def get_material_type(uploaded_file):
    content_type = uploaded_file.content_type

    if content_type.startswith('video/'):
        return 'video'

    if content_type.startswith('image/'):
        return 'image'

    return 'text'
        
def is_admin(user):
    return user.is_authenticated and user.is_staff


class AdminLoginView(LoginView):
    template_name = 'users/login.html'
    authentication_form = AdminLoginForm

    def get_success_url(self):
        return '/users/dashboard/'


class AdminLogoutView(LogoutView):
    pass


@login_required
@user_passes_test(is_admin)
def admin_register_user(request):
    if request.method == 'POST':
        form = AdminUserRegisterForm(request.POST)
        if form.is_valid():
            user = form.save()
            messages.success(request, f'Пользователь "{user.username}" успешно создан.')
            return redirect('home')
    else:
        form = AdminUserRegisterForm()

    return render(request, 'users/admin_register_user.html', {'form': form})

@login_required
@admin_required
def dashboard(request):
    partners_qs = PartnerRegistration.objects.filter(referral_user=request.user)

    total_partners = partners_qs.count()
    active_partners = partners_qs.filter(is_active_partner=True).count()
    inactive_partners = partners_qs.filter(is_active_partner=False).count()

    today = timezone.localdate()
    new_partners_month = partners_qs.filter(
        created_at__year=today.year,
        created_at__month=today.month
    ).count()
    new_partners_month = partners_qs.filter(
        created_at__year=today.year,
        created_at__month=today.month
    ).count()
    leads_qs = Lead.objects.filter(admin=request.user)

    total_leads = leads_qs.count()
    new_leads = leads_qs.filter(status='new').count()
    in_progress_leads = leads_qs.filter(status='in_progress').count()
    cancelled_leads = leads_qs.filter(status='cancelled').count()
    deal_leads = leads_qs.filter(status='deal').count()
    total_accrued = Accrual.objects.filter(
        admin=request.user
    ).aggregate(
        total=Sum('amount')
    )['total'] or 0

    deal_leads_qs = Lead.objects.filter(
    admin=request.user,
    status='deal'
)

    total_deal_amount = deal_leads_qs.aggregate(
        total=Coalesce(Sum('deal_amount'), Decimal('0'))
    )['total']

    total_partner_reward = deal_leads_qs.aggregate(
        total=Coalesce(Sum('partner_reward'), Decimal('0'))
    )['total']

    admin_profit = total_deal_amount - total_partner_reward

    lead_conversion = round((deal_leads / total_leads) * 100) if total_leads else 0
   
    paid_total = WithdrawalRequest.objects.filter(
        admin=request.user,
        status='paid'
    ).aggregate(
        total=Sum('amount')
    )['total'] or Decimal('0')

    pending_total = total_partner_reward - paid_total   
   
    withdrawal_requests_count = WithdrawalRequest.objects.filter(
        admin=request.user
    ).count()

    paid_withdrawal_requests_count = WithdrawalRequest.objects.filter(
        admin=request.user,
        status='paid'
    ).count()
    top_partners = (
    Lead.objects.filter(
        admin=request.user,
        status='deal'
    )
    .values(
        'partner__first_name',
        'partner__last_name'
    )
    .annotate(
        total_deals=Count('id')
    )
    .order_by('-total_deals')[:5]
    )
    top_partners_labels = [
    f"{item['partner__first_name']} {item['partner__last_name']}"
    for item in top_partners
]

    top_partners_values = [
        item['total_deals']
        for item in top_partners
    ]
    sales_by_month = (
    Lead.objects.filter(
        admin=request.user,
        status='deal'
    )
    .annotate(
        month=TruncMonth('created_at')
    )
    .values('month')
    .annotate(
        total_sales=Sum('deal_amount')
    )
    .order_by('month')
    )
    months_ru = {
        1: 'Янв',
        2: 'Фев',
        3: 'Мар',
        4: 'Апр',
        5: 'Май',
        6: 'Июн',
        7: 'Июл',
        8: 'Авг',
        9: 'Сен',
        10: 'Окт',
        11: 'Ноя',
        12: 'Дек',
    }

    sales_month_labels = [
        months_ru[item['month'].month]
        for item in sales_by_month
    ]

    sales_month_values = [
        float(item['total_sales']) for item in sales_by_month
    ]
    return render(request, 'users/dashboard.html', {
        'active_page': 'dashboard',
        'total_partners': total_partners,
        'active_partners': active_partners,
        'inactive_partners': inactive_partners,
        'new_partners_month': new_partners_month,
        'total_leads': total_leads,
        'new_leads': new_leads,
        'in_progress_leads': in_progress_leads,
        'cancelled_leads': cancelled_leads,
        'deal_leads': deal_leads,
        'lead_conversion': lead_conversion,
        'total_accrued': total_accrued,
        'admin_profit': admin_profit,
        'paid_total': paid_total,
        'pending_total': pending_total,
        'total_partner_reward': total_partner_reward,
        'withdrawal_requests_count': withdrawal_requests_count,
        'paid_withdrawal_requests_count': paid_withdrawal_requests_count,
        'top_partners_labels': top_partners_labels,
        'top_partners_values': top_partners_values,
        'sales_month_labels': sales_month_labels,
        'sales_month_values': sales_month_values,
    })
@login_required
@admin_required
def partners(request):  
    date_filter = request.GET.get('date_filter', 'all')
    partner_activity = request.GET.get('partner_activity', '')
    offer_id = request.GET.get('offer', '')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    search = request.GET.get('search', '')
    leads_count_subquery = Lead.objects.filter(
        partner=OuterRef('user')
    ).values('partner').annotate(
        total=Count('id')
    ).values('total')[:1]

    deals_count_subquery = Lead.objects.filter(
        partner=OuterRef('user'),
        status='deal'
    ).values('partner').annotate(
        total=Count('id')
    ).values('total')[:1]

    reward_subquery = Lead.objects.filter(
        partner=OuterRef('user'),
        status='deal'
    ).values('partner').annotate(
        total=Sum('partner_reward')
    ).values('total')[:1]

    paid_subquery = WithdrawalRequest.objects.filter(
        partner=OuterRef('user'),
        status='paid'
    ).values('partner').annotate(
        total=Sum('amount')
    ).values('total')[:1]
    
    base_queryset = PartnerRegistration.objects.filter(
    referral_user=request.user
    ).select_related(
        'status',
        'offer',
        'user'
    ).annotate(
        real_leads_count=Coalesce(
            Subquery(leads_count_subquery, output_field=IntegerField()),
            Value(0)
        ),
        real_deals_count=Coalesce(
            Subquery(deals_count_subquery, output_field=IntegerField()),
            Value(0)
        ),
        real_reward_amount=Coalesce(
            Subquery(reward_subquery, output_field=DecimalField()),
            Value(0),
            output_field=DecimalField()
        ),
        paid_amount=Coalesce(
            Subquery(paid_subquery, output_field=DecimalField()),
            Value(0),
            output_field=DecimalField()
        )
    ).annotate(
        real_balance=F('real_reward_amount') - F('paid_amount')
    ).order_by('-created_at')
    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    # Счетчики по вкладкам даты
    count_all = base_queryset.count()
    count_today = base_queryset.filter(created_at__date=today).count()
    count_yesterday = base_queryset.filter(created_at__date=yesterday).count()
    count_week = base_queryset.filter(created_at__date__gte=week_start).count()
    count_month = base_queryset.filter(created_at__date__gte=month_start).count()

    count_period = base_queryset
    if start_date:
        count_period = count_period.filter(created_at__date__gte=start_date)
    if end_date:
        count_period = count_period.filter(created_at__date__lte=end_date)
    count_period = count_period.count()

    partners_list = base_queryset

    if date_filter == 'today':
        partners_list = partners_list.filter(created_at__date=today)
    elif date_filter == 'yesterday':
        partners_list = partners_list.filter(created_at__date=yesterday)
    elif date_filter == 'week':
        partners_list = partners_list.filter(created_at__date__gte=week_start)
    elif date_filter == 'month':
        partners_list = partners_list.filter(created_at__date__gte=month_start)
    elif date_filter == 'period':
        if start_date:
            partners_list = partners_list.filter(created_at__date__gte=start_date)
        if end_date:
            partners_list = partners_list.filter(created_at__date__lte=end_date)

    # Счетчики активности в рамках уже выбранного периода/поиска/оффера
    filtered_for_activity = partners_list

    if offer_id:
        filtered_for_activity = filtered_for_activity.filter(offer_id=offer_id)

    if search:
        filtered_for_activity = filtered_for_activity.filter(
            Q(full_name__icontains=search) |
            Q(email__icontains=search) |
            Q(phone__icontains=search) |
            Q(company_name__icontains=search) |
            Q(company_short_name__icontains=search) |
            Q(company_full_name__icontains=search)
        )

    active_count = filtered_for_activity.filter(is_active_partner=True).count()
    inactive_count = filtered_for_activity.filter(is_active_partner=False).count()
    total_count = filtered_for_activity.count()

    # Применяем фильтры к итоговому списку
    if partner_activity == 'active':
        partners_list = partners_list.filter(is_active_partner=True)
    elif partner_activity == 'inactive':
        partners_list = partners_list.filter(is_active_partner=False)

    if offer_id:
        partners_list = partners_list.filter(offer_id=offer_id)

    if search:
        partners_list = partners_list.filter(
            Q(full_name__icontains=search) |
            Q(email__icontains=search) |
            Q(phone__icontains=search) |
            Q(company_name__icontains=search) |
            Q(company_short_name__icontains=search) |
            Q(company_full_name__icontains=search)
        )

    offer_count_filter = Q(
        partner_registrations__is_approved=True,
        partner_registrations__referral_user=request.user
    )

    if date_filter == 'today':
        offer_count_filter &= Q(partner_registrations__created_at__date=today)

    elif date_filter == 'yesterday':
        offer_count_filter &= Q(partner_registrations__created_at__date=yesterday)

    elif date_filter == 'week':
        offer_count_filter &= Q(partner_registrations__created_at__date__gte=week_start)

    elif date_filter == 'month':
        offer_count_filter &= Q(partner_registrations__created_at__date__gte=month_start)

    elif date_filter == 'period':
        if start_date:
            offer_count_filter &= Q(partner_registrations__created_at__date__gte=start_date)
        if end_date:
            offer_count_filter &= Q(partner_registrations__created_at__date__lte=end_date)

    user_offers = Offer.objects.filter(current_user=request.user).annotate(
        partners_count=Count(
            'partner_registrations',
            filter=offer_count_filter
        )
    ).order_by('title')
    partner_create_form = PartnerCreateByUserForm(user=request.user)

    

    return render(request, 'users/partners.html', {
        'partners_list': partners_list,
        'partner_create_form': partner_create_form,
        'active_page': 'partners',
        'date_filter': date_filter,
        'partner_activity': partner_activity,
        'selected_offer': offer_id,
        'start_date': start_date,
        'end_date': end_date,
        'search': search,
        'user_offers': user_offers,

        'count_all': count_all,
        'count_today': count_today,
        'count_yesterday': count_yesterday,
        'count_week': count_week,
        'count_month': count_month,
        'count_period': count_period,

        'active_count': active_count,
        'inactive_count': inactive_count,
        'total_count': total_count,
        
    })
@login_required
@admin_required
def referral_program(request):
    filter_type = request.GET.get('filter', 'all')
    search = request.GET.get('search', '').strip()
    all_offers = list(
        Offer.objects.filter(current_user=request.user).order_by('-created_at')
    )

    active_offers = [offer for offer in all_offers if offer.status == 'Активный']
    inactive_offers = [offer for offer in all_offers if offer.status == 'Неактивный']
    total_leads = Lead.objects.filter(
        admin=request.user
    ).count()

    total_deals = Lead.objects.filter(
        admin=request.user,
        status='deal'
    ).count()
    conversion_offers = [
        offer for offer in all_offers
        if Lead.objects.filter(
            admin=request.user,
            offer=offer,
            status='deal'
        ).exists()
    ]
    conversion_percent = round((total_deals / total_leads) * 100) if total_leads else 0
    if filter_type == 'active':
        offers = active_offers
    elif filter_type == 'inactive':
        offers = inactive_offers
    elif filter_type == 'conversion':
        offers = conversion_offers
    else:
        offers = all_offers
    if search:
        offers = [
            offer for offer in offers
            if search.lower() in offer.title.lower()
            or search.lower() in offer.description.lower()
            or search in str(offer.offer_id)
        ]
    for offer in offers:
        offer.deals_amount = Lead.objects.filter(
            admin=request.user,
            offer=offer,
            status='deal'
        ).aggregate(
            total=Sum('partner_reward')
        )['total'] or Decimal('0')

    create_form = OfferCreateForm()
    update_forms = {offer.id: OfferUpdateForm(instance=offer) for offer in offers}
    partner_statuses = PartnerStatus.objects.all().order_by('reward_percent')
    total_income = sum(offer.deals_amount for offer in offers)

    total_clicks = sum(
        offer.clicks_count for offer in offers
    )

    total_registrations = sum(
        offer.registrations_count for offer in offers
    )

    overall_conversion = round(
        (total_registrations / total_clicks) * 100
    ) if total_clicks else 0

    return render(request, 'users/referral_program.html', {
        'active_page': 'referral_program',
        'offers': offers,
        'form': create_form,
        'update_forms': update_forms,
        'current_filter': filter_type,
        'offers_count_all': len(all_offers),
        'offers_count_active': len(active_offers),
        'offers_count_inactive': len(inactive_offers),
        'offers_count_conversion': len(conversion_offers),
        'conversion_percent': conversion_percent,
        'partner_statuses': partner_statuses,
        'total_income': total_income,
        'total_clicks': total_clicks,
        'total_registrations': total_registrations,
        'overall_conversion': overall_conversion,
    })

@login_required
@admin_required
def finances(request):
    return render(request, 'users/finances.html', {
        'active_page': 'finances'
    })

@login_required
@admin_required
def profile(request):
    if request.method == 'POST':
        form = ProfileUpdateForm(
            request.POST,
            request.FILES,
            instance=request.user
        )
        if form.is_valid():
            user = form.save(commit=False)
            new_password = request.POST.get('new_password', '').strip()

            if new_password:
                user.set_password(new_password)

            if form.cleaned_data.get('delete_avatar'):
                if user.avatar:
                    user.avatar.delete(save=False)
                user.avatar = None

            user.save()
            if new_password:
                update_session_auth_hash(request, user)
            messages.success(request, 'Профиль успешно обновлен.')
            return redirect('profile')
    else:
        form = ProfileUpdateForm(instance=request.user)
    partners_count = PartnerRegistration.objects.filter(
        referral_user=request.user
    ).count()

    offers_count = Offer.objects.filter(
        current_user=request.user
    ).count()

    leads_count = Lead.objects.filter(
        admin=request.user
    ).count()

    deals_count = Lead.objects.filter(
        admin=request.user,
        status='deal'
    ).count()
    return render(request, 'users/profile.html', {
        'active_page': 'profile',
        'form': form,
        'partners_count': partners_count,
        'offers_count': offers_count,
        'leads_count': leads_count,
        'deals_count': deals_count,
    })
@login_required
@admin_required
def create_offer(request):
    if request.method == 'POST':
        form = OfferCreateForm(request.POST)
        files = request.FILES.getlist('promo_materials')

        if form.is_valid():
            offer = form.save(commit=False)
            offer.current_user = request.user
            offer.landing_page = 'https://partnetix.ru'
            offer.save()

            save_offer_files(offer, files)
            for promo_file in files:
                material_type = get_material_type(promo_file)

                MarketingMaterial.objects.create(
                    offer=offer,
                    title=promo_file.name,
                    description=f'Промоматериал для оффера: {offer.title}',
                    material_type=material_type,
                    file=promo_file,
                )
            messages.success(request, 'Оффер успешно добавлен.')
        else:
            messages.error(request, 'Проверьте форму оффера.')

    return redirect('referral_program')

@login_required
@admin_required
def update_offer(request, offer_id):
    offer = get_object_or_404(Offer, id=offer_id, current_user=request.user)

    if request.method == 'POST':
        form = OfferUpdateForm(request.POST, instance=offer)
        files = request.FILES.getlist('promo_materials')

        if form.is_valid():
            updated_offer = form.save(commit=False)
            updated_offer.current_user = request.user
            updated_offer.save()

            save_offer_files(updated_offer, files)

            messages.success(request, 'Оффер успешно обновлен.')
        else:
            messages.error(request, 'Проверьте данные формы.')

    return redirect('referral_program')
@login_required
@admin_required
def delete_offer_file(request, file_id):
    file_obj = get_object_or_404(
        OfferPromoMaterial,
        id=file_id,
        offer__current_user=request.user,
    )

    if request.method == 'POST':
        file_obj.file.delete(save=False)
        file_obj.delete()

        if request.headers.get('x-requested-with') == 'XMLHttpRequest':
            from django.http import JsonResponse
            return JsonResponse({'status': 'ok'})

        messages.success(request, 'Файл удален.')

    return redirect('referral_program')

@login_required
@admin_required
def create_partner_by_user(request):
    if request.method != 'POST':
        return redirect('partners')

    form = PartnerCreateByUserForm(request.POST, user=request.user)

    if form.is_valid():
        UserModel = get_user_model()
        default_status = PartnerStatus.objects.order_by('id').first()
        password = get_random_string(10)

        email = form.cleaned_data['email']
        partner_type = form.cleaned_data['partner_type']
        offer = form.cleaned_data['offer']

        user = UserModel.objects.create_user(
            username=email,
            email=email,
            password=password,
        )

        user.is_active = True
        user.referred_by = request.user

        if partner_type in ['individual', 'self_employed']:
            user.first_name = form.cleaned_data.get('full_name', '')
        elif partner_type == 'company':
            user.first_name = form.cleaned_data.get('contact_person_name', '')

        user.phone = form.cleaned_data.get('phone', '')

        if partner_type == 'self_employed':
            user.company = form.cleaned_data.get('company_name', '')
        elif partner_type == 'company':
            user.company = form.cleaned_data.get('company_short_name', '')

        user.save()

        registration = form.save(commit=False)
        registration.user = user
        registration.offer = offer
        registration.referral_user = request.user
        registration.status = default_status
        registration.is_approved = True
        registration.is_active_partner = True
        registration.save()
        partner_url = f'{settings.SITE_URL}/?referral={user.user_id}&offer={offer.offer_id}'

        PartnerLink.objects.get_or_create(
            partner=user,
            offer=offer,
            defaults={
                'title': 'Моя новая ссылка',
                'url': partner_url,
            }
        )
        PartnerLink.objects.get_or_create(
            partner=user,
            offer=offer,
            defaults={
                'title': 'Моя новая ссылка',
                'url': partner_url,
            }
        )

        messages.success(
            request,
            f'Партнер успешно добавлен. Временный пароль: {password}'
        )
    else:
        print(form.errors)
        messages.error(request, 'Проверьте форму добавления партнера.')

    return redirect('partners')



def partners_register(request):
    UserModel = get_user_model()
    referral_user = None
    offer = None
    inviter_name = ''
    ip_address = request.META.get('REMOTE_ADDR')
    if request.method == 'GET' and referral_user and offer:
        OfferVisit.objects.create(
            offer=offer,
            referral_user=referral_user,
            ip_address=ip_address,
            is_registered=False,
            partner_link=partner_link,
        )
    referral_id = request.GET.get('referral')
    offer_id = request.GET.get('offer')
    link_code = request.GET.get('link') or request.POST.get('link')

    partner_link = None
    if link_code:
        partner_link = PartnerLink.objects.filter(
            url__contains=f'link={link_code}'
        ).first()

    referral_user = None
    offer = None
    inviter_name = ''
    if referral_id:
        try:
            referral_user = User.objects.get(user_id=referral_id)
        except User.DoesNotExist:
            referral_user = None

    if offer_id:
        try:
            offer = Offer.objects.get(offer_id=offer_id)
        except Offer.DoesNotExist:
            offer = None
    if not offer:
        default_offers = Offer.objects.filter(is_default=True)

        for default_offer in default_offers:
            if default_offer.status == 'Активный':
                offer = default_offer
                break

        partner_link = None

        if link_code:
            partner_link = PartnerLink.objects.filter(
                url__contains=f'link={link_code}'
            ).first()
    if request.method == 'GET' and referral_user and offer:
        ip_address = request.META.get('REMOTE_ADDR')

    OfferVisit.objects.create(
        offer=offer,
        referral_user=referral_user,
        ip_address=ip_address,
        is_registered=False,
        partner_link=partner_link,
    )
    if referral_user:
        inviter_name = f'{referral_user.first_name} {referral_user.last_name}'.strip()
        if not inviter_name:
            inviter_name = referral_user.username

    default_status = PartnerStatus.objects.order_by('id').first()
    partner_type = request.GET.get('type', 'individual')

    if request.method == 'POST':
        partner_type = request.POST.get('partner_type', 'individual')
        form = PartnerRegistrationForm(request.POST)

        form.fields['status_display'].initial = default_status.name if default_status else ''
        form.fields['inviter_name'].initial = inviter_name

        if form.is_valid():
            email = form.cleaned_data['email']
            password = form.cleaned_data['password1']

            user = UserModel.objects.create_user(
                username=email,
                email=email,
                password=password,
            )

            user.is_active = False

            if partner_type in ['individual', 'self_employed']:
                user.first_name = form.cleaned_data.get('full_name', '')
            elif partner_type == 'company':
                user.first_name = form.cleaned_data.get('contact_person_name', '')

            user.phone = form.cleaned_data.get('phone', '')

            if partner_type == 'self_employed':
                user.company = form.cleaned_data.get('company_name', '')
            elif partner_type == 'company':
                user.company = form.cleaned_data.get('company_short_name', '')

            if referral_user:
                user.referred_by = referral_user

            user.save()

            registration = form.save(commit=False)
            registration.partner_type = partner_type
            registration.offer = offer
            registration.referral_user = referral_user
            registration.status = default_status
            registration.user = user
            registration.is_approved = False
            registration.partner_link = partner_link
            registration.save()

            messages.success(
                request,
                'Заявка отправлена. Вход будет доступен после подтверждения модератором.'
            )
            return redirect('/partner/login/')
    else:
        form = PartnerRegistrationForm(initial={
            'partner_type': partner_type,
            'status_display': default_status.name if default_status else '',
            'inviter_name': inviter_name,
        })

    return render(request, 'users/partners_register.html', {
        'form': form,
        'partner_type': partner_type,
        'offer': offer,
        'referral_user': referral_user,
        'referral_id': referral_id,
        'offer_id': offer_id,
        'inviter_name': inviter_name,
        'default_status': default_status,
    })
@login_required
@admin_required
def update_partner_status(request, status_id):
    if request.method != 'POST':
        return redirect('/users/referral-program/?filter=statuses')
    status = get_object_or_404(PartnerStatus, id=status_id)

    reward_percent = request.POST.get('reward_percent')

    try:
        reward_percent = int(reward_percent)
    except (TypeError, ValueError):
        reward_percent = None

    if reward_percent is not None and 0 <= reward_percent <= 100:
        status.reward_percent = reward_percent
        status.save()
        messages.success(request, 'Процент статуса обновлён.')
    else:
        messages.error(request, 'Введите процент от 0 до 100.')

    return redirect('/users/referral-program/?filter=statuses')

@login_required
@admin_required
def leads(request):
    return render(request, 'users/leads.html', {
        'active_page': 'leads'
    })
@login_required
@admin_required
def reports(request):
    report_generated = False
    generate_report = request.GET.get('generate_report') == '1'
    partner_id = request.GET.get('partner', '')
    report_type = request.GET.get('report_type', 'all')
    date_filter = request.GET.get('date_filter', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')
    period_text = ''
    selected_partner_reg = None
    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)
    if date_filter == 'today':
        period_text = f'за {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'yesterday':
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'

    elif date_filter == 'week':
        period_text = f'за период с {week_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'month':
        period_text = f'за период с {month_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'
    elif date_filter == 'period':
        if start_date and end_date:
            period_text = f'за период с {start_date} по {end_date}'

        elif start_date:
            period_text = f'с {start_date}'

        elif end_date:
            period_text = f'по {end_date}'    
    total_leads = 0
    deal_leads = 0
    total_amount = 0
    total_reward = 0
    partners_for_report = PartnerRegistration.objects.filter(
        referral_user=request.user,
        is_approved=True,
        user__isnull=False
    ).select_related('user').order_by('user__first_name')

    leads_report = None
    selected_partner = None

    if partner_id and generate_report:
        selected_partner_reg = get_object_or_404(
            PartnerRegistration,
            referral_user=request.user,
            user_id=partner_id,
            is_approved=True,
            user__isnull=False
        )

        selected_partner = selected_partner_reg.user
        leads_report = Lead.objects.filter(
            admin=request.user,
            partner=selected_partner
        ).select_related(
            'partner',
            'offer',
            'tariff',
            'client'
        ).order_by('-created_at')
        if report_type == 'deals':
            leads_report = leads_report.filter(status='deal')
        elif report_type == 'no_deals':
            leads_report = leads_report.exclude(status='deal')
        if date_filter == 'today':
            leads_report = leads_report.filter(created_at__date=today)
        elif date_filter == 'yesterday':
            leads_report = leads_report.filter(created_at__date=yesterday)
        elif date_filter == 'week':
            leads_report = leads_report.filter(created_at__date__gte=week_start)
        elif date_filter == 'month':
            leads_report = leads_report.filter(created_at__date__gte=month_start)
        elif date_filter == 'period':
            start_date_text = ''
            end_date_text = ''

            if start_date:
                start_date_text = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d.%m.%Y')

            if end_date:
                end_date_text = datetime.strptime(end_date, '%Y-%m-%d').strftime('%d.%m.%Y')

            if start_date and end_date:
                period_text = f'за период с {start_date_text} по {end_date_text}'

            elif start_date:
                period_text = f'с {start_date_text}'

            elif end_date:
                period_text = f'по {end_date_text}'
        total_leads = leads_report.count()
        deal_leads = leads_report.filter(status='deal').count()
        total_amount = sum(lead.deal_amount or 0 for lead in leads_report)
        total_reward = sum(lead.partner_reward or 0 for lead in leads_report)
        paginator = Paginator(leads_report, 10)
        page_number = request.GET.get('page')
        leads_report = paginator.get_page(page_number)
        report_generated = leads_report.paginator.count > 0
    return render(request, 'users/reports.html', {
        'active_page': 'reports',
        'partners_for_report': partners_for_report,
        'selected_partner': partner_id,
        'selected_partner_obj': selected_partner,
        'leads_report': leads_report,
        'date_filter': date_filter,
        'start_date': start_date,
        'end_date': end_date,
        'report_type': report_type,
        'total_leads': total_leads,
        'deal_leads': deal_leads,
        'total_amount': total_amount,
        'total_reward': total_reward,
        'selected_partner_reg': selected_partner_reg,
        'period_text': period_text,
        'report_generated': report_generated,
    })
@login_required
@admin_required
def education(request):
    return render(request, 'users/education.html', {
        'active_page': 'education'
    })


@login_required
@admin_required
def materials(request):
    return render(request, 'users/materials.html', {
        'active_page': 'materials'
    })

@login_required
@admin_required
def partner_chat(request):
    return render(request, 'users/partner_chat.html', {
        'active_page': 'partner_chat'
    })
@login_required
@admin_required
def leads(request):
    leads_list = Lead.objects.filter(
        admin=request.user
    ).select_related(
        'partner',
        'offer'
    ).order_by('-created_at')
    paid_accruals = Accrual.objects.filter(
    lead=OuterRef('pk'),
    payout_status='paid'
    )

    leads_list = leads_list.annotate(
        is_paid_out=Exists(paid_accruals)
    )
    date_filter = request.GET.get('date_filter', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    partner_id = request.GET.get('partner', '')
    offer_id = request.GET.get('offer', '')
    status = request.GET.get('status', '')
    search = request.GET.get('search', '').strip()
    tabs_queryset = Lead.objects.filter(
        admin=request.user
    )
    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)
    if partner_id:
        tabs_queryset = tabs_queryset.filter(partner_id=partner_id)

    if offer_id:
        tabs_queryset = tabs_queryset.filter(offer_id=offer_id)

    if status:
        tabs_queryset = tabs_queryset.filter(status=status)
    leads_list = tabs_queryset  
    if date_filter == 'today':
        leads_list = tabs_queryset.filter(created_at__date=today)
    elif date_filter == 'yesterday':
        leads_list = tabs_queryset.filter(created_at__date=yesterday)
    elif date_filter == 'week':
        leads_list = tabs_queryset.filter(created_at__date__gte=week_start)
    elif date_filter == 'month':
        leads_list = tabs_queryset.filter(created_at__date__gte=month_start)
    elif date_filter == 'period':
        if start_date:
            leads_list = tabs_queryset.filter(created_at__date__gte=start_date)
        if end_date:
            leads_list = tabs_queryset.filter(created_at__date__lte=end_date)
   
    if partner_id:
        leads_list = leads_list.filter(partner_id=partner_id)

    if offer_id:
        leads_list = leads_list.filter(offer_id=offer_id)

    if status:
        leads_list = leads_list.filter(status=status)

    if search:
        status_map = {
            'новый': 'new',
            'новый лид': 'new',
            'в работе': 'in_progress',
            'отмена': 'cancelled',
            'сделка': 'deal',
        }

        search_lower = search.lower()
        matched_status = status_map.get(search_lower)

        search_query = (
            Q(title__icontains=search) |
            Q(partner__first_name__icontains=search) |
            Q(partner__email__icontains=search) |
            Q(offer__title__icontains=search) |
            Q(client__first_name__icontains=search) |
            Q(client__email__icontains=search) |
            Q(client__phone__icontains=search)
        )

        if matched_status:
            search_query |= Q(status=matched_status)

        leads_list = leads_list.filter(search_query)
    partners_for_lead = PartnerRegistration.objects.filter(
            referral_user=request.user,
            is_approved=True,
            user__isnull=False,
        ).select_related(
            'user',
            'offer',
            'status'
        ).annotate(
            partner_leads_total=Count('user__leads')
        )

    tariffs = Tariff.objects.filter(
        is_active=True
    ).order_by('period', 'price')
  
  
    offers = Offer.objects.filter(current_user=request.user).order_by('title')

    offer_counts = dict(
        Lead.objects.filter(
            admin=request.user,
            offer__isnull=False
        )
        .values_list('offer_id')
        .annotate(count=Count('id'))
    )

    for offer in offers:
        offer.leads_count = offer_counts.get(offer.id, 0)
  
  
  
    base_queryset = Lead.objects.filter(admin=request.user)
    total_deal_amount = base_queryset.aggregate(
    total=Sum('deal_amount')
    )['total'] or 0
    total_partner_reward = base_queryset.aggregate(
        total=Sum('partner_reward')
    )['total'] or 0
    partners_count = partners_for_lead.count()
    offers_count = offers.count()

    count_status_all = leads_list.count()
    count_status_new = leads_list.filter(status='new').count()
    count_status_in_progress = leads_list.filter(status='in_progress').count()
    count_status_cancelled = leads_list.filter(status='cancelled').count()
    count_status_deal = leads_list.filter(status='deal').count()
    page_number = request.GET.get('page', 1)
    summary_queryset = leads_list

    summary_total_leads = summary_queryset.count()

    summary_new_leads = summary_queryset.filter(
        status='new'
    ).count()

    summary_in_progress_leads = summary_queryset.filter(
        status='in_progress'
    ).count()

    summary_cancelled_leads = summary_queryset.filter(
        status='cancelled'
    ).count()

    summary_deal_leads = summary_queryset.filter(
        status='deal'
    ).count()

    summary_total_deal_amount = summary_queryset.aggregate(
        total=Sum('deal_amount')
    )['total'] or 0

    summary_total_partner_reward = summary_queryset.aggregate(
        total=Sum('partner_reward')
    )['total'] or 0

    summary_conversion = round(
        (summary_deal_leads / summary_total_leads) * 100
    ) if summary_total_leads else 0
    paginator = Paginator(leads_list, 25)  # 25 лидов на страницу
    leads_list = paginator.get_page(page_number)
    query_params = request.GET.copy()

    if 'page' in query_params:
        query_params.pop('page')

    pagination_query = query_params.urlencode()
    total_leads = base_queryset.count()
    new_leads = base_queryset.filter(status='new').count()
    in_progress_leads = base_queryset.filter(status='in_progress').count()
    cancelled_leads = base_queryset.filter(status='cancelled').count()
    deal_leads = base_queryset.filter(status='deal').count()

    conversion = round((deal_leads / total_leads) * 100) if total_leads else 0
    return render(request, 'users/leads.html', {
        'active_page': 'leads',
        'leads_list': leads_list,
        'partners_for_lead': partners_for_lead,
        'tariffs': tariffs,
        'offers': offers,
        'date_filter': date_filter,
        'start_date': start_date,
        'end_date': end_date,
        'count_all': base_queryset.count(),
        'count_all': tabs_queryset.count(),
        'count_today': tabs_queryset.filter(created_at__date=today).count(),
        'count_yesterday': tabs_queryset.filter(created_at__date=yesterday).count(),
        'count_week': tabs_queryset.filter(created_at__date__gte=week_start).count(),
        'count_month': tabs_queryset.filter(created_at__date__gte=month_start).count(),
        'count_period': tabs_queryset.count(),
        'partners_count': partners_count,
        'offers_count': offers_count,
        'count_status_all': count_status_all,
        'count_status_new': count_status_new,
        'count_status_in_progress': count_status_in_progress,
        'count_status_cancelled': count_status_cancelled,
        'count_status_deal': count_status_deal,
        'selected_partner': partner_id,
        'selected_offer': offer_id,
        'selected_status': status,
        'search': search,
        'pagination_query': pagination_query,
        'total_leads': total_leads,
        'new_leads': new_leads,
        'in_progress_leads': in_progress_leads,
        'cancelled_leads': cancelled_leads,
        'deal_leads': deal_leads,
        'conversion': conversion,
        'total_deal_amount': total_deal_amount,
        'total_partner_reward': total_partner_reward,
        'summary_total_leads': summary_total_leads,
        'summary_new_leads': summary_new_leads,
        'summary_in_progress_leads': summary_in_progress_leads,
        'summary_cancelled_leads': summary_cancelled_leads,
        'summary_deal_leads': summary_deal_leads,
        'summary_total_deal_amount': summary_total_deal_amount,
        'summary_total_partner_reward': summary_total_partner_reward,
        'summary_conversion': summary_conversion,
    })

def client_register(request):

    UserModel = get_user_model()

    referral_id = request.GET.get('referral') or request.POST.get('referral')
    offer_id = request.GET.get('offer') or request.POST.get('offer')
    link_code = request.GET.get('link') or request.POST.get('link')
    partner = None
    offer = None

    if referral_id:
        partner = User.objects.filter(user_id=referral_id).first()

    if offer_id:
        offer = Offer.objects.filter(offer_id=offer_id).first()
    partner_link = None

    if link_code:
        partner_link = PartnerLink.objects.filter(
            url__contains=f'link={link_code}'
        ).first()
    if request.method == 'GET' and partner and offer:
        OfferVisit.objects.create(
            offer=offer,
            referral_user=partner,
            ip_address=request.META.get('REMOTE_ADDR'),
            is_registered=False,
            partner_link=partner_link,
        )    
    if request.method == 'POST':
        form = ClientRegistrationForm(request.POST)

        if form.is_valid():
            email = form.cleaned_data['email']
            full_name = form.cleaned_data['full_name']
            phone = form.cleaned_data.get('phone', '')
            company = form.cleaned_data.get('company', '')
            password = form.cleaned_data['password1']

            user = UserModel.objects.create_user(
                username=email,
                email=email,
                password=password,
            )

            user.first_name = full_name
            user.phone = phone
            user.company = company

            if partner:
                user.referred_by = partner

            user.save()
            if partner:
                Lead.objects.create(
                    title='Новый лид',
                    deal_amount=0,
                    partner_reward=0,
                    partner=partner,
                    client=user,
                    status='new',
                    offer=offer,
                    admin=partner.referred_by,
                    partner_link=partner_link,
                )

            messages.success(request, 'Регистрация успешно завершена.')
            return redirect('login')
    else:
        form = ClientRegistrationForm()

    return render(request, 'users/client_register.html', {
        'form': form,
        'referral_id': referral_id,
        'offer_id': offer_id,
        'partner': partner,
        'offer': offer,
        'link_code': link_code,
    })

@login_required
@admin_required
def create_lead(request):
    if request.method != 'POST':
        return redirect('leads')

    partner_id = request.POST.get('partner')
    tariff_id = request.POST.get('tariff')
    title = request.POST.get('title', '').strip() or 'Новый лид'
    status = request.POST.get('status', 'new')

    client_name = request.POST.get('client_name', '').strip()
    client_email = request.POST.get('client_email', '').strip()
    client_phone = request.POST.get('client_phone', '').strip()
    client_company = request.POST.get('client_company', '').strip()

    partner = get_object_or_404(User, id=partner_id)
    tariff = None
    if tariff_id:
        tariff = get_object_or_404(Tariff, id=tariff_id)

    registration = PartnerRegistration.objects.filter(
        user=partner,
        referral_user=request.user,
        is_approved=True,
    ).select_related('offer', 'status').first()

    if not registration:
        messages.error(request, 'Партнер не найден или не привязан к вам.')
        return redirect('leads')

    offer = registration.offer

    client = User.objects.filter(
    Q(email=client_email) | Q(username=client_email)
    ).first()

    if client:
        messages.error(request, 'Клиент с таким email уже существует.')
        return redirect('leads')
    with transaction.atomic():
        password = get_random_string(10)

        client = User.objects.create_user(
            username=client_email,
            email=client_email,
            password=password,
        )

        client.first_name = client_name
        client.phone = client_phone
        client.company = client_company
        client.referred_by = partner
        client.save()

        deal_amount = tariff.price if tariff else Decimal('0')
        partner_reward = Decimal('0')

        if status == 'deal' and offer and offer.status == 'Активный':
            if offer.payout_type == 'fixed':
                partner_reward = offer.reward

            elif offer.payout_type == 'partner_status' and registration.status:
                partner_reward = (
                    deal_amount * registration.status.reward_percent / Decimal('100')
                )
        else:
            partner_reward = Decimal('0')

        lead=Lead.objects.create(
            title=title,
            status=status,
            tariff=tariff,
            deal_amount=deal_amount,
            partner_reward=partner_reward,
            partner=partner,
            client=client,
            offer=offer,
            admin=request.user,            
        )
        if lead.status == 'deal' and lead.partner_reward > 0:
            if not Accrual.objects.filter(lead=lead).exists():
                Accrual.objects.create(
                    amount=lead.partner_reward,
                    lead=lead,
                    partner=lead.partner,
                    offer=lead.offer,
                    admin=request.user,
                    
                )
    if password:
        messages.success(request, f'Лид добавлен. Временный пароль клиента: {password}')
    else:
        messages.success(request, 'Лид добавлен.')
    return redirect('leads')

@login_required
@admin_required
def update_lead(request, lead_id):
    print('UPDATE_LEAD CALLED')
    lead = get_object_or_404(Lead, id=lead_id, admin=request.user)

    print('LEAD ID:', lead.id)

    print(
        'ACCRUALS:',
        list(
            Accrual.objects.filter(
                lead=lead
            ).values(
                'id',
                'payout_status'
            )
        )
    )

    print(
        'WITHDRAWALS:',
        list(
            WithdrawalRequest.objects.filter(
                accruals__lead=lead
            ).values(
                'id',
                'status'
            )
        )
    )


    if Accrual.objects.filter(
        lead=lead
    ).filter(
        models.Q(payout_status='paid') |
        models.Q(withdrawal_requests__status='paid')
    ).exists():
        messages.error(
            request,
            'Этот лид нельзя изменить, так как по нему уже была произведена выплата.'
        )
        return redirect('leads')

    if request.method != 'POST':
        return redirect('leads')

    title = request.POST.get('title', '').strip() or 'Новый лид'
    status = request.POST.get('status', 'new')
    tariff_id = request.POST.get('tariff')
    partner_id = request.POST.get('partner')
    offer_id = request.POST.get('offer')
    created_at = request.POST.get('created_at')

    if partner_id:
        lead.partner = get_object_or_404(User, id=partner_id)

    lead.title = title
    lead.status = status

    tariff = None
    if status == 'deal' and tariff_id:
        tariff = get_object_or_404(Tariff, id=tariff_id)

    lead.tariff = tariff
    lead.deal_amount = tariff.price if tariff else Decimal('0')

    if offer_id:
        offer = get_object_or_404(Offer, id=offer_id)
    else:
        registration = PartnerRegistration.objects.filter(
            user=lead.partner,
            referral_user=request.user,
            is_approved=True,
        ).select_related('offer').first()
    

    lead.offer = offer

    registration = PartnerRegistration.objects.filter(
        user=lead.partner,
        referral_user=request.user,
        is_approved=True,
    ).select_related('status').first()
    
    offer = registration.offer if registration else None
    partner_reward = Decimal('0')

    if status == 'deal' and offer and offer.status == 'Активный':
        if offer.payout_type == 'fixed':
            partner_reward = offer.reward

        elif (
            offer.payout_type == 'partner_status'
            and registration
            and registration.status
            and tariff
        ):
            partner_reward = (
                lead.deal_amount * registration.status.reward_percent / Decimal('100')
            )

    if status != 'deal':
        tariff = None
        lead.tariff = None
        lead.deal_amount = Decimal('0')
        partner_reward = Decimal('0')
    lead.partner_reward = partner_reward

    if created_at:
        lead.created_at = datetime.strptime(created_at, '%Y-%m-%d')

    if lead.client:
        client_email = request.POST.get('client_email', '').strip()

        if client_email:
            lead.client.first_name = request.POST.get('client_name', '').strip()
            lead.client.email = client_email
            lead.client.username = client_email
            lead.client.phone = request.POST.get('client_phone', '').strip()
            lead.client.company = request.POST.get('client_company', '').strip()
            lead.client.save()

    lead.save()
    print('STATUS:', lead.status)
    print('REWARD:', lead.partner_reward)
    print('ACCRUALS:', Accrual.objects.filter(lead=lead).count())

    if lead.status == 'deal' and lead.partner_reward > 0:
        accrual, created = Accrual.objects.get_or_create(
            lead=lead,
            defaults={
                'amount': lead.partner_reward,
                'partner': lead.partner,
                'offer': lead.offer,
                'admin': request.user,
            }
        )

        if not created:
            accrual.amount = lead.partner_reward
            accrual.partner = lead.partner
            accrual.offer = lead.offer
            accrual.admin = request.user
            accrual.save()
    else:
        Accrual.objects.filter(lead=lead).delete()

    messages.success(request, 'Лид успешно обновлен.')
    return redirect('leads')

@login_required
@admin_required
def client_tariffs(request):
    tariffs = Tariff.objects.filter(
        is_active=True
    ).order_by('period', 'price')

    referral_id = request.GET.get('referral', '')
    offer_id = request.GET.get('offer', '')
    link_code = request.GET.get('link', '')
    subscription = Subscription.objects.filter(
        user=request.user,
        status='active'
    ).select_related('tariff').first()
    subscription = Subscription.objects.filter(
        user=request.user,
        status='active'
    ).select_related('tariff').first()
    active_tariff_id = subscription.tariff_id if subscription and subscription.tariff else None
    return render(request, 'users/client_tariffs.html', {
        'tariffs': tariffs,
        'referral_id': referral_id,
        'offer_id': offer_id,
        'link_code': link_code,
        'active_page': 'tariffs',
        'active_tariff_id': active_tariff_id,
        'subscription': subscription,
    })
@login_required
@admin_required
def payouts(request):
    status_filter = request.GET.get('status', 'all')

    base_queryset = WithdrawalRequest.objects.filter(
        admin=request.user
    )

    payout_requests = base_queryset.select_related(
        'partner'
    ).prefetch_related(
        'accruals',
        'accruals__lead',
        'accruals__offer'
    ).order_by('-created_at')

    if status_filter == 'pending':
        payout_requests = payout_requests.filter(
            status__in=['new', 'processing']
        )

    elif status_filter == 'paid':
        payout_requests = payout_requests.filter(
            status='paid'
        )

    return render(request, 'users/payouts.html', {
        'active_page': 'payouts',
        'payout_requests': payout_requests,
        'status_filter': status_filter,

        'all_count': base_queryset.count(),
        'pending_count': base_queryset.filter(
            status__in=['new', 'processing']
        ).count(),
        'paid_count': base_queryset.filter(
            status='paid'
        ).count(),
    })
def decrypt_card_number(encrypted_card_number):
    f = Fernet(settings.CARD_ENCRYPTION_KEY.encode())
    return f.decrypt(encrypted_card_number.encode()).decode()

@login_required
@admin_required
def payout_card_number(request, payout_id):
    withdrawal = get_object_or_404(
        WithdrawalRequest,
        id=payout_id,
        admin=request.user
    )

    if not withdrawal.encrypted_card_number:
        return JsonResponse({
            'success': False,
            'message': 'Номер карты отсутствует.'
        })

    try:
        card_number = decrypt_card_number(
            withdrawal.encrypted_card_number
        )

        return JsonResponse({
            'success': True,
            'card_number': card_number
        })

    except Exception as e:
        return JsonResponse({
            'success': False,
            'message': f'Ошибка расшифровки: {str(e)}'
        })
@login_required
@admin_required
def mark_payout_paid(request, payout_id):
    if request.method != 'POST':
        return redirect('payouts')

    payout = get_object_or_404(
        WithdrawalRequest,
        id=payout_id,
        admin=request.user
    )

    payout.status = 'paid'
    payout.processed_at = timezone.now()
    payout.save()

    payout.accruals.update(
        payout_status='paid'
    )

    messages.success(request, 'Выплата отмечена как завершённая.')
    return redirect('payouts')    
@login_required
@admin_required
def export_reports_csv(request):
    partner_id = request.GET.get('partner', '')
    report_type = request.GET.get('report_type', 'all')

    date_filter = request.GET.get('date_filter', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    leads = Lead.objects.none()

    if partner_id:
        partner = get_object_or_404(
            User,
            id=partner_id,
            referred_by=request.user
        )

        leads = Lead.objects.filter(
            admin=request.user,
            partner=partner
        ).select_related(
            'client',
            'offer'
        ).order_by('-created_at')

        if report_type == 'deals':
            leads = leads.filter(status='deal')
        elif report_type == 'no_deals':
            leads = leads.exclude(status='deal')

        if date_filter == 'today':
            leads = leads.filter(created_at__date=today)
        elif date_filter == 'yesterday':
            leads = leads.filter(created_at__date=yesterday)
        elif date_filter == 'week':
            leads = leads.filter(created_at__date__gte=week_start)
        elif date_filter == 'month':
            leads = leads.filter(created_at__date__gte=month_start)
        elif date_filter == 'period':
            if start_date:
                leads = leads.filter(created_at__date__gte=start_date)
            if end_date:
                leads = leads.filter(created_at__date__lte=end_date)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="report.csv"'

    response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
    response['Content-Disposition'] = 'attachment; filename="report.csv"'
    response.write('\ufeff')

    writer = csv.writer(response, delimiter=';')

    writer.writerow([
        'Дата',
        'Лид',
        'Клиент',
        'Оффер',
        'Статус',
        'Сумма сделки',
        'Вознаграждение'
    ])

    for lead in leads:
        writer.writerow([
            lead.created_at.strftime('%d.%m.%Y %H:%M'),
            lead.title,
            lead.client.first_name if lead.client else '',
            lead.offer.title if lead.offer else '',
            lead.get_status_display(),
            lead.deal_amount,
            lead.partner_reward,
        ])

    return response
def set_table_width(table, width):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn('w:tblW'))

    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)

    tbl_w.set(qn('w:w'), str(width))
    tbl_w.set(qn('w:type'), 'dxa')

def format_money(value):
    return f'{int(value):,}'.replace(',', ' ')

def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()

    tc_w = tc_pr.find(qn('w:tcW'))
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)

    tc_w.set(qn('w:w'), str(width))
    tc_w.set(qn('w:type'), 'dxa')
def set_table_width(table, width):
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn('w:tblW'))
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)

    tbl_w.set(qn('w:w'), str(width))
    tbl_w.set(qn('w:type'), 'dxa')

    tbl_layout = OxmlElement('w:tblLayout')
    tbl_layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(tbl_layout)
    
@login_required
@admin_required
def export_reports_word(request):
    partner_id = request.GET.get('partner', '')
    partner_reg = PartnerRegistration.objects.filter(
        referral_user=request.user,
        user_id=partner_id,
        is_approved=True
    ).first()

    partner_name = 'партнёру'

    if partner_reg:
        if partner_reg.partner_type == 'company':
            partner_name = partner_reg.company_short_name or partner_reg.company_full_name
        else:
            partner_name = partner_reg.user.first_name or partner_reg.user.email
    leads = Lead.objects.filter(
        admin=request.user,
        partner_id=partner_id
    ).select_related(
        'client',
        'offer'
    ).order_by('-created_at')
    report_type = request.GET.get('report_type', 'all')
    date_filter = request.GET.get('date_filter', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    if report_type == 'deals':
        leads = leads.filter(status='deal')
    elif report_type == 'no_deals':
        leads = leads.exclude(status='deal')

    if date_filter == 'today':
        leads = leads.filter(created_at__date=today)
    elif date_filter == 'yesterday':
        leads = leads.filter(created_at__date=yesterday)
    elif date_filter == 'week':
        leads = leads.filter(created_at__date__gte=week_start)
    elif date_filter == 'month':
        leads = leads.filter(created_at__date__gte=month_start)
    elif date_filter == 'period':
        if start_date:
            leads = leads.filter(created_at__date__gte=start_date)
        if end_date:
            leads = leads.filter(created_at__date__lte=end_date)
    document = Document()
    section = document.sections[0]

    section.orientation = WD_ORIENT.LANDSCAPE

    section.page_width = Mm(297)
    section.page_height = Mm(210)

    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)

    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    title = document.add_heading(f'Отчёт по партнёру: {partner_name}', level=1)
    title.runs[0].font.size = Pt(16)

    table = document.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.autofit = False
    set_table_width(table, 15500)
    headers = [
        'Дата',
        'Лид',
        'Клиент',
        'Оффер',
        'Статус',
        'Сумма',
        'Вознаграждение'
    ]

    widths = [
        Inches(1.5),
        Inches(1.6),
        Inches(2.2),
        Inches(2.5),
        Inches(1.2),
        Inches(1.4),
        Inches(1.5),
    ]

    hdr_cells = table.rows[0].cells

    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].width = widths[i]

        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for lead in leads:
        row = table.add_row().cells

        values = [
            lead.created_at.strftime('%d.%m.%Y %H:%M'),
            lead.title or '',
            lead.client.first_name if lead.client else '',
            lead.offer.title if lead.offer else '',
            lead.get_status_display(),
            f'{lead.deal_amount or 0} ₽',
            f'{lead.partner_reward or 0} ₽',
        ]

        for i, value in enumerate(values):
            row[i].text = str(value)
            row[i].width = widths[i]

            for paragraph in row[i].paragraphs:
                for run in paragraph.runs:
                        run.font.size = Pt(9)
    document.add_paragraph('')

    summary_table = document.add_table(rows=1, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    summary_table.autofit = False
    set_table_width(summary_table, 15500)
    total_leads = leads.count()
    deal_leads = leads.filter(status='deal').count()

    total_amount = sum(lead.deal_amount or 0 for lead in leads)
    total_reward = sum(lead.partner_reward or 0 for lead in leads)
    summary_cells = summary_table.rows[0].cells


    summary_widths = [3300, 3300, 4500, 4400]

    for i, cell in enumerate(summary_cells):
        set_cell_width(cell, summary_widths[i])
    summary_cells[0].text = f'Лидов: {total_leads}'
    summary_cells[1].text = f'Сделок: {deal_leads}'
    summary_cells[2].text = f'Сумма сделок: {format_money(total_amount)} ₽'
    summary_cells[3].text = f'Вознаграждение: {format_money(total_reward)} ₽'

    for row in summary_table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()

            borders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                borders.append(border)

            tc_pr.append(borders)

            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        document.add_paragraph('')

    signatures = document.add_table(rows=1, cols=3)
    signatures.autofit = False

    signatures.columns[0].width = Inches(2.5)
    signatures.columns[1].width = Inches(2.5)
    signatures.columns[2].width = Inches(2)

    signatures.cell(0, 0).text = 'Заказчик __________________'
    signatures.cell(0, 1).text = 'Исполнитель __________________'
    signatures.cell(0, 2).text = timezone.now().strftime('Дата: %d.%m.%Y')

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    response['Content-Disposition'] = 'attachment; filename=report.docx'
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
    document.add_paragraph('')
    section = document.sections[0]

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]

    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = footer_paragraph.add_run('Страница ')

    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.text = 'PAGE'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)
    document.save(response)

    return response
def add_page_number(canvas, doc):
    canvas.setFont('Arial', 8)

    page_num = canvas.getPageNumber()

    canvas.drawRightString(
        285 * mm,
        10 * mm,
        f'Страница {page_num}'
    )
@login_required
@admin_required
def export_reports_pdf(request):
    partner_id = request.GET.get('partner', '')
    report_type = request.GET.get('report_type', 'all')
    date_filter = request.GET.get('date_filter', 'all')
    start_date = request.GET.get('start_date', '')
    end_date = request.GET.get('end_date', '')

    today = timezone.localdate()
    yesterday = today - timedelta(days=1)
    week_start = today - timedelta(days=7)
    month_start = today - timedelta(days=30)

    partner_reg = PartnerRegistration.objects.filter(
        referral_user=request.user,
        user_id=partner_id,
        is_approved=True
    ).select_related('user').first()

    partner_name = ''

    if partner_reg:
        if partner_reg.partner_type == 'company':
            partner_name = partner_reg.company_short_name or partner_reg.company_full_name
        else:
            partner_name = partner_reg.user.first_name or partner_reg.user.email

    leads = Lead.objects.filter(
        admin=request.user,
        partner_id=partner_id
    ).select_related(
        'client',
        'offer'
    ).order_by('-created_at')

    if report_type == 'deals':
        leads = leads.filter(status='deal')
    elif report_type == 'no_deals':
        leads = leads.exclude(status='deal')

    if date_filter == 'today':
        leads = leads.filter(created_at__date=today)
        period_text = f'за {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'yesterday':
        leads = leads.filter(created_at__date=yesterday)
        period_text = f'за {yesterday.strftime("%d.%m.%Y")}'

    elif date_filter == 'week':
        leads = leads.filter(created_at__date__gte=week_start)
        period_text = f'за период с {week_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'month':
        leads = leads.filter(created_at__date__gte=month_start)
        period_text = f'за период с {month_start.strftime("%d.%m.%Y")} по {today.strftime("%d.%m.%Y")}'

    elif date_filter == 'period':
        if start_date:
            leads = leads.filter(created_at__date__gte=start_date)
        if end_date:
            leads = leads.filter(created_at__date__lte=end_date)

        start_text = start_date
        end_text = end_date

        if start_date:
            start_text = datetime.strptime(start_date, '%Y-%m-%d').strftime('%d.%m.%Y')
        if end_date:
            end_text = datetime.strptime(end_date, '%Y-%m-%d').strftime('%d.%m.%Y')

        period_text = f'за период с {start_text} по {end_text}'

    else:
        period_text = ''

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="report.pdf"'

    pdfmetrics.registerFont(TTFont('Arial', r'C:\Windows\Fonts\arial.ttf'))
    pdfmetrics.registerFont(TTFont('Arial-Bold', r'C:\Windows\Fonts\arialbd.ttf'))

    registerFontFamily(
        'Arial',
        normal='Arial',
        bold='Arial-Bold',
        italic='Arial',
        boldItalic='Arial-Bold'
    )

    doc = SimpleDocTemplate(
        response,
        pagesize=landscape(A4),
        rightMargin=10 * mm,
        leftMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )

    styles = getSampleStyleSheet()

    styles['Normal'].fontName = 'Arial'
    styles['Normal'].fontSize = 8
    styles['Normal'].leading = 10

    styles['Title'].fontName = 'Arial'
    styles['Title'].fontSize = 18
    styles['Title'].leading = 24
    styles['Title'].alignment = TA_CENTER

    period_style = styles['Normal'].clone('PeriodStyle')
    period_style.alignment = TA_CENTER
    period_style.fontName = 'Arial'
    period_style.fontSize = 10
    period_style.leading = 12

    elements = []

    elements.append(Paragraph(
        f'<b>Отчёт по партнёру: {partner_name}</b>',
        styles['Title']
    ))

    if period_text:
        elements.append(Paragraph(period_text, period_style))
        elements.append(Spacer(1, 12))
    else:
        elements.append(Spacer(1, 8))

    data = [[
        'Дата',
        'Лид',
        'Клиент',
        'Оффер',
        'Статус',
        'Сумма сделки',
        'Вознаграждение',
    ]]

    for lead in leads:
        client = ''
        if lead.client:
            client = f'{lead.client.first_name}<br/>{lead.client.email}'

        data.append([
            lead.created_at.strftime('%d.%m.%Y %H:%M'),
            lead.title or '',
            Paragraph(client, styles['Normal']),
            Paragraph(lead.offer.title if lead.offer else '', styles['Normal']),
            lead.get_status_display(),
            f'{format_money(lead.deal_amount or 0)} ₽',
            f'{format_money(lead.partner_reward or 0)} ₽',
        ])

    table_widths = [
        32 * mm,
        28 * mm,
        45 * mm,
        45 * mm,
        28 * mm,
        30 * mm,
        35 * mm,
    ]

    report_width = sum(table_widths)

    table = Table(data, colWidths=table_widths)

    table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f3f4f6')),
        ('GRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d1d5db')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),

        ('FONTNAME', (0, 0), (-1, 0), 'Arial-Bold'),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#111827')),

        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
    ]))

    elements.append(table)
    elements.append(Spacer(1, 12))

    total_leads = leads.count()
    deal_leads = leads.filter(status='deal').count()
    total_amount = sum(lead.deal_amount or 0 for lead in leads)
    total_reward = sum(lead.partner_reward or 0 for lead in leads)

    summary = Table(
        [[
            f'Лидов: {total_leads}',
            f'Сделок: {deal_leads}',
            f'Сумма сделок: {format_money(total_amount)} ₽',
            f'Вознаграждение: {format_money(total_reward)} ₽',
        ]],
        colWidths=[
            report_width / 4,
            report_width / 4,
            report_width / 4,
            report_width / 4,
        ]
    )

    summary.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Arial-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#f3f4f6')),
        ('BOX', (0, 0), (-1, -1), 0.25, colors.HexColor('#d1d5db')),
        ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.HexColor('#d1d5db')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
    ]))

    elements.append(summary)
    elements.append(Spacer(1, 18))

    signatures = Table(
        [[
            'Заказчик ____________________',
            'Исполнитель ____________________',
            f'Дата: {today.strftime("%d.%m.%Y")}',
        ]],
        colWidths=[
            report_width / 3,
            report_width / 3,
            report_width / 3,
        ]
    )

    signatures.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Arial'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (0, 0), 'LEFT'),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
        ('ALIGN', (2, 0), (2, 0), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))

    elements.append(signatures)

    doc.build(
        elements,
        onFirstPage=add_page_number,
        onLaterPages=add_page_number
    )

    return response

@login_required
@admin_required
def education(request):
    videos = EducationVideo.objects.order_by('-created_at')
    documents = EducationDocument.objects.order_by('-created_at')

    context = {
        'videos': videos,
        'documents': documents,
        'active_page': 'education',
    }

    return render(request, 'users/education.html', context)

@login_required
@admin_required
def add_education_video(request):
    if request.method != 'POST':
        return redirect('education')

    title = request.POST.get('title', '').strip()
    description = request.POST.get('description', '').strip()
    duration = request.POST.get('duration', '').strip()
    video_file = request.FILES.get('video_file')

    if title and video_file:

        temp_video = VideoFileClip(video_file.temporary_file_path())

        total_seconds = int(temp_video.duration)
        temp_video.close()
        minutes = total_seconds // 60
        seconds = total_seconds % 60

        duration = f"{minutes}:{seconds:02d}"

        EducationVideo.objects.create(
            title=title,
            description=description,
            duration=duration,
            video_file=video_file,
        )

        messages.success(request, 'Видео успешно добавлено.')
    else:
        messages.error(request, 'Заполните название и выберите файл видео.')

    return redirect('education')

@login_required
@admin_required
def delete_education_video(request, video_id):
    video = get_object_or_404(EducationVideo, id=video_id)

    if request.method == 'POST':
        video.video_file.delete(save=False)
        video.delete()
        messages.success(request, 'Видео удалено.')

    return redirect('education')


@login_required
@admin_required
def replace_education_video(request, video_id):
    video = get_object_or_404(EducationVideo, id=video_id)

    if request.method != 'POST':
        return redirect('education')

    new_file = request.FILES.get('video_file')

    if not new_file:
        messages.error(request, 'Выберите новый файл видео.')
        return redirect('education')

    old_file_path = video.video_file.path if video.video_file else None

    video.video_file = new_file

    temp_video = VideoFileClip(new_file.temporary_file_path())
    total_seconds = int(temp_video.duration)
    temp_video.close()

    minutes = total_seconds // 60
    seconds = total_seconds % 60
    video.duration = f'{minutes}:{seconds:02d}'

    video.save()

    if old_file_path:
        try:
            if os.path.exists(old_file_path):
                os.remove(old_file_path)
        except PermissionError:
            messages.warning(
                request,
                'Видео заменено, но старый файл сейчас занят и не был удалён.'
            )
            return redirect('education')

    messages.success(request, 'Видео заменено.')
    return redirect('education')
@login_required
def add_education_document(request):
    if request.method != 'POST':
        return redirect('education')

    title = request.POST.get('title', '').strip()
    description = request.POST.get('description', '').strip()
    document_file = request.FILES.get('document_file')

    if title and document_file:
        EducationDocument.objects.create(
            title=title,
            description=description,
            document_file=document_file,
        )
        messages.success(request, 'Материал успешно добавлен.')
    else:
        messages.error(request, 'Заполните название и выберите файл.')

    return redirect('education')


@login_required
def delete_education_document(request, doc_id):
    document = get_object_or_404(EducationDocument, id=doc_id)

    if request.method == 'POST':
        document.document_file.delete(save=False)
        document.delete()
        messages.success(request, 'Материал удалён.')

    return redirect('education')

@login_required
def materials(request):
    materials = MarketingMaterial.objects.select_related('offer').order_by('-created_at')

    videos = materials.filter(material_type='video')
    images = materials.filter(material_type='image')
    texts = materials.filter(material_type='text')
    context = {
        'materials': materials,
        'videos': videos,
        'images': images,
        'texts': texts,
        'active_page': 'materials',
        'offers': Offer.objects.all(),
    }

    return render(request, 'users/materials.html', context)

@login_required
@admin_required
def create_material(request):
    if request.method == 'POST':
        MarketingMaterial.objects.create(
            offer_id=request.POST.get('offer'),
            title=request.POST.get('title'),
            description=request.POST.get('description'),
            material_type=request.POST.get('material_type'),
            file=request.FILES.get('file'),
            text_content=request.POST.get('text_content'),
        )

    return redirect('materials')

@login_required
@admin_required
def delete_material(request, material_id):
    if request.method == 'POST':
        material = get_object_or_404(MarketingMaterial, id=material_id)
        OfferPromoMaterial.objects.filter(
        offer=material.offer,
        file=material.file.name
        ).delete()

        material.delete()

    return redirect('materials')

@login_required
@admin_required
def admin_chat(request):
    topics = (
        ChatTopic.objects.filter(admin=request.user)
        .select_related(
            'partner',
            'partner__partner_registration'
        )
        .prefetch_related('messages')
        .annotate(
            unread_count=Count(
                'messages',
                filter=Q(
                    messages__is_read=False
                ) & ~Q(
                    messages__sender=request.user
                )
            )
        )
        .order_by('-updated_at')
    )
    base_topics = topics
    all_count = base_topics.count()
    open_count = base_topics.filter(status='open').count()
    closed_count = base_topics.filter(status='closed').count()
    unread_count = base_topics.filter(unread_count__gt=0).count()
    partners = (
        User.objects.filter(
            role='partner',
            referred_by=request.user
        )
        .annotate(
            chat_topics_count=Count('partner_chat_topics')
        )
        .order_by('first_name')
    )
    current_filter = request.GET.get('filter', 'all')
    if current_filter == 'open':
        topics = topics.filter(status='open')
    elif current_filter == 'closed':
        topics = topics.filter(status='closed')
    elif current_filter == 'unread':
        topics = topics.filter(unread_count__gt=0)
    all_topics = ChatTopic.objects.filter(
        admin=request.user
    ).select_related('partner').order_by('-updated_at')    
    return render(request, 'users/admin_chat.html', {
        'active_page': 'admin_chat',
        'topics': topics,
        'partners': partners,
        'current_filter': current_filter,
        'all_count': all_count,
        'open_count': open_count,
        'closed_count': closed_count,
        'unread_count': unread_count,
        'all_topics': all_topics,

    })
@login_required
@admin_required
def admin_chat_topic(request, topic_id):
    topics = (
    ChatTopic.objects.filter(admin=request.user)
        .select_related(
            'partner',
            'partner__partner_registration'
        )
        .prefetch_related('messages')
        .annotate(
            unread_count=Count(
                'messages',
                filter=Q(
                    messages__is_read=False
                ) & ~Q(
                    messages__sender=request.user
                )
            )
        )
        .order_by('-updated_at')
    )
    base_topics = topics

    all_count = base_topics.count()
    open_count = base_topics.filter(status='open').count()
    closed_count = base_topics.filter(status='closed').count()
    unread_count = base_topics.filter(unread_count__gt=0).count()

    current_filter = request.GET.get('filter', 'all')

    if current_filter == 'open':
        topics = topics.filter(status='open')

    elif current_filter == 'closed':
        topics = topics.filter(status='closed')

    elif current_filter == 'unread':
        topics = topics.filter(unread_count__gt=0)
    for topic in topics:
        last_message = topic.messages.order_by('-created_at').first()

        topic.last_message_obj = last_message

        if last_message:
            topic.last_sender_label = (
                'Партнёр'
                if last_message.sender.role == 'partner'
                else 'Администратор'
            )
        else:
            topic.last_sender_label = ''
    current_topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )
    current_topic.messages.filter(
        is_read=False
    ).exclude(
        sender=request.user
    ).update(
        is_read=True
    )
    messages = list(
        current_topic.messages
        .select_related('sender')
        .order_by('-created_at')[:100]
    )

    messages.reverse()
    partners = (
        User.objects.filter(
            role='partner',
            referred_by=request.user
        )
        .annotate(
            chat_topics_count=Count('partner_chat_topics')
        )
        .order_by('first_name')
    )

    if current_filter == 'open':
        topics = topics.filter(status='open')
    elif current_filter == 'closed':
        topics = topics.filter(status='closed')
    elif current_filter == 'unread':
        topics = topics.filter(unread_count__gt=0)
    all_topics = ChatTopic.objects.filter(
        admin=request.user
    ).select_related('partner').order_by('-updated_at') 
    partner_reg = getattr(
    current_topic.partner,
    'partner_registration',
    None
)
    if partner_reg and partner_reg.partner_type == 'company':
        current_topic.partner_display_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or current_topic.partner.email
        )
    else:
        current_topic.partner_display_name = (
            f'{current_topic.partner.first_name} '
            f'{current_topic.partner.last_name}'
        ).strip() or current_topic.partner.emai     
    return render(request, 'users/admin_chat.html', {
        'active_page': 'admin_chat',
        'topics': topics,
        'current_topic': current_topic,
        'messages': messages,
        'partners': partners,
        'current_filter': current_filter,
        'all_count': all_count,
        'open_count': open_count,
        'closed_count': closed_count,
        'unread_count': unread_count,
        'all_topics': all_topics
    })
@login_required
@admin_required
def admin_chat_send_message(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    if request.method == 'POST':
        if topic.status == 'closed':
            return JsonResponse({
                'success': False,
                'error': 'Тема закрыта'
            })
        text = request.POST.get('text', '').strip()

        if text:
            message = ChatMessage.objects.create(
                topic=topic,
                sender=request.user,
                text=text,
                is_read=False
            )
            topic.updated_at = timezone.now()
            topic.save(update_fields=['updated_at'])
            topic.status = 'open'
            topic.save(update_fields=['status', 'updated_at'])
            request.user.typing_updated_at = None
            request.user.save(update_fields=['typing_updated_at'])
            sender_name = ''
            partner_registration = getattr(
                request.user,
                'partner_registration',
                None
            )

            if (
                partner_registration
                and partner_registration.partner_type == 'company'
            ):
                sender_name = (
                    partner_registration.company_short_name
                    or partner_registration.company_full_name
                    or request.user.email
                )
            else:
                sender_name = (
                    f'{request.user.first_name} {request.user.last_name}'
                ).strip() or request.user.email

            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                channel_layer = get_channel_layer()

                async_to_sync(channel_layer.group_send)(
                    f'user_notifications_{topic.partner.id}',
                    {
                        'type': 'topic_message_created',
                        'topic_id': topic.id,
                        'text': message.text,
                        'time': message.created_at.strftime('%H:%M'),
                        'sender_name': sender_name,
                        'sender_email': request.user.email,
                        'sender_role': 'Администратор',
                    }
                )
                return JsonResponse({
                    'success': True,
                    'id': message.id,
                    'text': message.text,
                    'time': message.created_at.strftime('%H:%M'),
                    'is_own': True,
                    'sender_name': sender_name,
                    'sender_email': request.user.email,
                    'sender_role': 'Администратор',
                })

    return redirect('admin_chat_topic', topic_id=topic.id)

@login_required
def chat_send_image(request, topic_id):
    topic = get_object_or_404(ChatTopic, id=topic_id)

    if request.user != topic.admin and request.user != topic.partner:
        return JsonResponse({
            'success': False,
            'error': 'Нет доступа'
        }, status=403)

    if request.method != 'POST':
        return JsonResponse({'success': False}, status=405)

    image = request.FILES.get('image')

    if not image:
        return JsonResponse({
            'success': False,
            'error': 'Изображение не передано'
        }, status=400)

    message = ChatMessage.objects.create(
        topic=topic,
        sender=request.user,
        message_type='image',
        image_file=image
    )

    return JsonResponse({
        'success': True,
        'id': message.id,
        'message_type': message.message_type,
        'image_url': message.image_file.url,
        'time': message.created_at.strftime('%H:%M'),
        'is_own': True,
        'is_read': False,
    })

@login_required
@admin_required
def admin_chat_toggle_topic_status(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    if request.method == 'POST':
        if topic.status == 'closed':
            topic.status = 'open'
        else:
            topic.status = 'closed'

        topic.save(update_fields=['status', 'updated_at'])

    return redirect('admin_chat_topic', topic_id=topic.id)
@login_required
@admin_required
def admin_chat_create_topic(request):
    if request.method == 'POST':
        partner_id = request.POST.get('partner')
        title = request.POST.get('title', '').strip()

        if partner_id and title:
            partner = get_object_or_404(
                User,
                id=partner_id,
                role='partner',
                referred_by=request.user
            )

            topic = ChatTopic.objects.create(
                title=title,
                partner=partner,
                admin=request.user,
                created_by=request.user,
                status='open',
                priority='normal',
                category='other'
            )
            channel_layer = get_channel_layer()

            async_to_sync(channel_layer.group_send)(
                f'user_notifications_{partner.id}',
                {
                    'type': 'topic_created',
                    'topic_id': topic.id,
                    'title': topic.title,
                    'url': reverse('partner_chat_topic', args=[topic.id]),
                    'status_display': topic.get_status_display(),
                    'created_by': 'Администратор',
                }
            )
            if request.headers.get('x-requested-with') == 'XMLHttpRequest':
                return JsonResponse({
                    'success': True,
                    'topic_id': topic.id,
                    'title': topic.title,
                    'url': reverse('admin_chat_topic', args=[topic.id]),
                    'partner_name': (
                        partner.partner_registration.company_short_name
                        if hasattr(partner, 'partner_registration')
                        and partner.partner_registration.partner_type == 'company'
                        else f'{partner.first_name} {partner.last_name}'
                    ),
                    'partner_email': partner.email,
                    'status': topic.status,
                    'status_display': topic.get_status_display(),
                })

            return redirect('admin_chat_topic', topic_id=topic.id)

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({
            'success': False,
            'error': 'Не удалось создать тему'
        })

    return redirect('admin_chat')

@login_required
@admin_required
def admin_chat_messages(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    after_id = request.GET.get('after_id', 0)

    try:
        after_id = int(after_id)
    except (TypeError, ValueError):
        after_id = 0

    messages = (
        topic.messages
        .select_related('sender')
        .filter(id__gt=after_id)
        .order_by('created_at')
    )
    data = []

    for message in messages:
        data.append({
            'id': message.id,
            'text': message.text,
            'time': message.created_at.strftime('%H:%M'),
            'is_read': message.is_read,
            'is_own': message.sender_id == request.user.id,
            'message_type': message.message_type,
            'audio_url': message.audio_file.url if message.audio_file else None,
            'image_url': message.image_file.url if message.image_file else None,
            'video_note_url': message.video_note.url if message.video_note else None,
            'file_url': message.file.url if message.file else None,
            'file_name': message.file.name.split('/')[-1] if message.file else None,
            'avatar_url': message.sender.avatar.url if message.sender.avatar else '/static/img/default-avatar.png',
        })
    typing = topic.partner.is_typing
    return JsonResponse({
        'messages': data,
        'typing': typing
    })

@login_required
@admin_required
def admin_chat_typing(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    if request.method == 'POST':
        request.user.typing_updated_at = timezone.now()
        request.user.save(update_fields=['typing_updated_at'])

    return JsonResponse({'success': True})

@login_required
@admin_required
def admin_chat_send_voice(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    if request.method == 'POST':
        audio = request.FILES.get('audio')

        if audio:
            message = ChatMessage.objects.create(
                topic=topic,
                sender=request.user,
                message_type='voice',
                audio_file=audio,
                text=''
            )

            topic.updated_at = timezone.now()
            topic.status = 'open'
            topic.save(update_fields=['updated_at', 'status'])

            return JsonResponse({
                'success': True,
                'id': message.id,
                'message_type': message.message_type,
                'audio_url': message.audio_file.url,
                'time': message.created_at.strftime('%H:%M'),
                'is_own': True,
            })

    return JsonResponse({
        'success': False
    })

@login_required
def user_heartbeat(request):
    request.user.last_activity = timezone.now()
    request.user.save(update_fields=['last_activity'])

    return JsonResponse({'success': True})

@login_required
@admin_required
def admin_chat_partners_statuses(request):
    partners = User.objects.filter(
        role='partner',
        referred_by=request.user
    )

    data = []

    for partner in partners:
        data.append({
            'id': partner.id,
            'is_typing': partner.is_typing,
            'is_online': partner.is_online,
            'last_seen_text': partner.last_seen_text,
        })

    return JsonResponse({
        'partners': data
    })
@login_required
@admin_required
def chat_read_statuses(request, topic_id):

    topic = get_object_or_404(ChatTopic, id=topic_id)

    messages = topic.messages.filter(
        sender=request.user
    )

    data = []

    for message in messages:
        data.append({
            'id': message.id,
            'is_read': message.is_read,
        })

    return JsonResponse({
        'messages': data
    })
@login_required
@admin_required
def admin_chat_mark_read(request, topic_id):

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    if request.method == 'POST':

        topic.messages.filter(
            is_read=False
        ).exclude(
            sender=request.user
        ).update(
            is_read=True
        )

    return JsonResponse({
        'success': True
    })

@login_required
def chat_send_file(request, topic_id):
    topic = get_object_or_404(ChatTopic, id=topic_id)

    if request.user != topic.admin and request.user != topic.partner:
        return JsonResponse({
            'success': False,
            'error': 'Нет доступа'
        }, status=403)

    if request.method != 'POST':
        return JsonResponse({'success': False}, status=405)

    uploaded_file = request.FILES.get('file')
    max_file_size = 100 * 1024 * 1024

    if uploaded_file.size > max_file_size:
        return JsonResponse({
            'success': False,
            'error': 'Файл слишком большой. Максимальный размер — 100 МБ.'
        }, status=400)
    if not uploaded_file:
        return JsonResponse({
            'success': False,
            'error': 'Файл не передан'
        }, status=400)

    message = ChatMessage.objects.create(
        topic=topic,
        sender=request.user,
        message_type='file',
        file=uploaded_file
    )

    file_url = message.file.url
    file_name = os.path.basename(message.file.name)

    return JsonResponse({
        'success': True,
        'id': message.id,
        'message_type': message.message_type,
        'file_url': file_url,
        'file_name': file_name,
        'time': message.created_at.strftime('%H:%M'),
        'is_own': True,
        'is_read': False,
    })
@login_required
@admin_required
def admin_chat_send_video_note(request, topic_id):
    if request.method != 'POST':
        return JsonResponse({'success': False})

    topic = get_object_or_404(
        ChatTopic,
        id=topic_id
    )

    video = request.FILES.get('video')

    if not video:
        return JsonResponse({'success': False})

    message = ChatMessage.objects.create(
        topic=topic,
        sender=request.user,
        message_type='video_note',
        video_note=video
    )

    return JsonResponse({
        'success': True,
        'id': message.id,
        'message_type': message.message_type,
        'video_note_url': message.video_note.url,
        'time': message.created_at.strftime('%H:%M'),
        'is_own': True,
        'is_read': False,
    })
import tempfile
from django.http import JsonResponse
from django.contrib.auth.decorators import login_required
from openai import OpenAI


@login_required
def chat_transcribe_voice(request, topic_id):
    if request.method != 'POST':
        return JsonResponse({
            'success': False,
            'error': 'Метод не поддерживается'
        })

    audio = request.FILES.get('audio')

    if not audio:
        return JsonResponse({
            'success': False,
            'error': 'Аудио не получено'
        })

    try:
        client = OpenAI()

        with tempfile.NamedTemporaryFile(
            suffix='.webm',
            delete=False
        ) as temp_audio:

            for chunk in audio.chunks():
                temp_audio.write(chunk)

            temp_audio_path = temp_audio.name

        with open(temp_audio_path, 'rb') as audio_file:
            transcription = client.audio.transcriptions.create(
                model='whisper-1',
                file=audio_file,
                language='ru'
            )

        return JsonResponse({
            'success': True,
            'text': transcription.text
        })

    except Exception as error:
        return JsonResponse({
            'success': False,
            'error': str(error)
        })

@login_required
@admin_required
def admin_chat_topic_data(request, topic_id):
    topic = get_object_or_404(
        ChatTopic,
        id=topic_id,
        admin=request.user
    )

    topic.messages.filter(
        is_read=False
    ).exclude(
        sender=request.user
    ).update(
        is_read=True
    )

    messages = (
        topic.messages
        .select_related('sender')
        .order_by('created_at')
    )

    partner_reg = getattr(
        topic.partner,
        'partner_registration',
        None
    )

    if partner_reg and partner_reg.partner_type == 'company':
        partner_name = (
            partner_reg.company_short_name
            or partner_reg.company_full_name
            or topic.partner.email
        )
    else:
        partner_name = (
            f'{topic.partner.first_name} {topic.partner.last_name}'
        ).strip() or topic.partner.email

    def serialize_message(message):
        return {
            'id': message.id,
            'text': message.text or '',
            'time': message.created_at.strftime('%H:%M'),
            'is_own': message.sender_id == request.user.id,
            'message_type': message.message_type,
            'is_read': message.is_read,

            'audio_url': message.audio_file.url if message.audio_file else '',
            'image_url': message.image_file.url if message.image_file else '',
            'file_url': message.file.url if message.file else '',
            'file_name': message.file_name if message.file else '',
            'video_note_url': message.video_note.url if message.video_note else '',

            'avatar_url': (
                message.sender.avatar.url
                if message.sender.avatar
                else ''
            ),
        }

    return JsonResponse({
        'topic': {
            'id': topic.id,
            'title': topic.title,
            'status': topic.status,
            'status_display': topic.get_status_display(),
            'partner_name': partner_name,
            'url': reverse('admin_chat_topic', args=[topic.id]),
        },
        'messages': [
            serialize_message(message)
            for message in messages
        ]
    })
@login_required
@admin_required
def partner_detail(request, partner_id):
    partner = get_object_or_404(
        User.objects.select_related('partner_registration'),
        id=partner_id
    )

    leads = Lead.objects.filter(
        partner=partner
    ).select_related(
        'offer',
        'client'
    ).order_by('-created_at')

    leads_count = leads.count()
    deals_count = leads.filter(status='deal').count()

    total_reward = (
        leads.filter(status='deal')
        .aggregate(total=Sum('partner_reward'))['total'] or 0
    )
    partner_registration = partner.partner_registration
    payment_details, created = PartnerPaymentDetails.objects.get_or_create(
    partner=partner
)
    if request.method == 'POST':
        print('PARTNER SAVE POST')
        print('POST DATA:', request.POST)
        print(request.POST)
        if partner_registration.partner_type == 'company':
            partner.email = request.POST.get('email', '').strip()
            partner.phone = request.POST.get('phone', '').strip()
            partner.save(update_fields=['email', 'phone'])

            partner_registration.company_full_name = request.POST.get(
                'company_full_name',
                ''
            ).strip()

            partner_registration.company_short_name = request.POST.get(
                'company_short_name',
                ''
            ).strip()

            partner_registration.legal_address = request.POST.get(
                'legal_address',
                ''
            ).strip()

            partner_registration.postal_address = request.POST.get(
                'postal_address',
                ''
            ).strip()

            partner_registration.contact_person_name = request.POST.get(
                'contact_person_name',
                ''
            ).strip()
            old_status_id = partner_registration.status_id
            status_id = request.POST.get('status')

            if status_id:
                partner_registration.status_id = status_id
            else:
                partner_registration.status = None

            partner_registration.save()
            if old_status_id != partner_registration.status_id:
                recalculate_partner_rewards(
                    partner,
                    partner_registration
                )            
            new_status_id = partner_registration.status_id

            if old_status_id != new_status_id and partner_registration.status:
                paid_lead_ids = Accrual.objects.filter(
                    lead__isnull=False
                ).filter(
                    models.Q(payout_status='paid') |
                    models.Q(withdrawal_requests__status='paid')
                ).values_list('lead_id', flat=True)

                leads_to_update = Lead.objects.filter(
                    partner=partner,
                    offer__payout_type='partner_status'
                ).exclude(
                    id__in=paid_lead_ids
                ).select_related('tariff')
                for lead in leads_to_update:
                    if lead.tariff:
                        lead.deal_amount = lead.tariff.price
                        lead.partner_reward = (
                            lead.deal_amount *
                            partner_registration.status.reward_percent /
                            Decimal('100')
                        )
                        lead.save(update_fields=[
                            'deal_amount',
                            'partner_reward'
                        ])

                        Accrual.objects.filter(
                            lead=lead
                        ).update(
                            amount=lead.partner_reward
                        )
            return JsonResponse({
                'success': True,
                'message': 'Изменения успешно сохранены'
            })

        elif partner_registration.partner_type == 'individual':
            partner.email = request.POST.get('email', '').strip()
            partner.phone = request.POST.get('phone', '').strip()
            partner.save(update_fields=['email', 'phone'])

            partner_registration.full_name = request.POST.get(
                'full_name',
                ''
            ).strip()

            partner_registration.activity_type = request.POST.get(
                'activity_type',
                ''
            ).strip()
            old_status_id = partner_registration.status_id
            status_id = request.POST.get('status')

            if status_id:
                partner_registration.status_id = status_id
            else:
                partner_registration.status = None

            partner_registration.save()
            if old_status_id != partner_registration.status_id:
                recalculate_partner_rewards(
                    partner,
                    partner_registration
                )
            return JsonResponse({
                'success': True,
                'message': 'Изменения успешно сохранены'
            })
        elif partner_registration.partner_type == 'self_employed':
            partner.email = request.POST.get('email', '').strip()
            partner.phone = request.POST.get('phone', '').strip()
            partner.save(update_fields=['email', 'phone'])

            partner_registration.full_name = request.POST.get(
                'full_name',
                ''
            ).strip()

            partner_registration.activity_type = request.POST.get(
                'activity_type',
                ''
            ).strip()

            partner_registration.company_name = request.POST.get(
                'company_name',
                ''
            ).strip()

            partner_registration.inn = request.POST.get(
                'inn',
                ''
            ).strip()
            old_status_id = partner_registration.status_id
            status_id = request.POST.get('status')

            if status_id:
                partner_registration.status_id = status_id
            else:
                partner_registration.status = None

            partner_registration.save()
            if old_status_id != partner_registration.status_id:
                recalculate_partner_rewards(
                    partner,
                    partner_registration
                )
            return JsonResponse({
                'success': True,
                'message': 'Изменения успешно сохранены'
            })
        return JsonResponse({
            'success': False,
            'message': 'Этот тип партнёра пока не обрабатывается'
        }, status=400)
    partner_statuses = PartnerStatus.objects.all()
    offers = partner_registration.offers.all()
    for offer in offers:
        offer.leads_count = Lead.objects.filter(
            partner=partner,
            offer=offer
        ).count()

        offer.deals_count = Lead.objects.filter(
            partner=partner,
            offer=offer,
            status='deal'
        ).count()

        offer.deals_amount = (
            Lead.objects.filter(
                partner=partner,
                offer=offer,
                status='deal'
            ).aggregate(total=Sum('partner_reward'))['total'] or 0
        )
    create_form = OfferCreateForm()
    all_offers = Offer.objects.exclude(
    id__in=partner_registration.offers.values_list(
        'id',
        flat=True
    )
    ).order_by('title')
    return render(request, 'users/partner_detail.html', {
        'active_page': 'partners',
        'partner': partner,
        'leads': leads[:5],
        'leads_count': leads_count,
        'deals_count': deals_count,
        'total_reward': total_reward,
        'partner_registration': partner_registration,
        'partner_statuses': partner_statuses,
        'offers': offers,
        'all_offers': all_offers,
        'create_form': create_form,
        'payment_details': payment_details,
    })

@login_required
@admin_required
def create_partner_offer(request, partner_id):
    if request.method != 'POST':
        return redirect('partner_detail', partner_id=partner_id)

    partner = get_object_or_404(
        User,
        id=partner_id
    )

    partner_registration = get_object_or_404(
        PartnerRegistration,
        user=partner
    )
    offer_mode = request.POST.get('offer_mode', 'new')

    if offer_mode == 'existing':
        existing_offer_id = request.POST.get('existing_offer_id')

        offer = get_object_or_404(
            Offer,
            id=existing_offer_id
        )

        partner_registration.offers.add(offer)

        if not partner_registration.offer:
            partner_registration.offer = offer
            partner_registration.save(update_fields=['offer'])

        landing_page = offer.landing_page or 'https://partnetix.ru'

        partner_url = (
            f'{landing_page}'
            f'/?referral={partner.user_id}'
            f'&offer={offer.offer_id}'
        )

        PartnerLink.objects.get_or_create(
            partner=partner,
            offer=offer,
            defaults={
                'title': 'Моя новая ссылка',
                'url': partner_url,
            }
        )

        messages.success(request, 'Существующий оффер подключён партнёру.')

        return redirect('partner_detail', partner_id=partner.id)
    title = request.POST.get('title', '').strip()
    description = request.POST.get('description', '').strip()
    payout_type = request.POST.get('payout_type', 'fixed')
    reward = request.POST.get('reward') or 0
    activity_start = request.POST.get('activity_start')
    activity_end = request.POST.get('activity_end')
    landing_page = request.POST.get('landing_page') or 'https://partnetix.ru'

    offer = Offer.objects.create(
        title=title,
        description=description,
        payout_type=payout_type,
        reward=reward,
        activity_start=activity_start,
        activity_end=activity_end,
        landing_page=landing_page,
        current_user=partner
    )

    partner_registration.offers.add(offer)

    if not partner_registration.offer:
        partner_registration.offer = offer
        partner_registration.save(update_fields=['offer'])
        partner_url = (
            f'{landing_page}'
            f'/?referral={partner.user_id}'
            f'&offer={offer.offer_id}'
        )
    partner_url = (
        f'{landing_page}'
        f'/?referral={partner.user_id}'
        f'&offer={offer.offer_id}'
    )
    PartnerLink.objects.get_or_create(
        partner=partner,
        offer=offer,
        defaults={
            'title': 'Моя новая ссылка',
            'url': partner_url,
        }
    )

    return redirect('partner_detail', partner_id=partner.id)

@login_required
@admin_required
def update_partner_offer(request, partner_id, offer_id):
    partner = get_object_or_404(User, id=partner_id)

    partner_registration = get_object_or_404(
        PartnerRegistration,
        user=partner
    )
    offer = get_object_or_404(
        partner_registration.offers,
        id=offer_id
    )

    if request.method == 'POST':
        old_payout_type = offer.payout_type
        old_reward = offer.reward
        offer.title = request.POST.get('title', '').strip()
        offer.description = request.POST.get('description', '').strip()
        offer.payout_type = request.POST.get('payout_type', 'fixed')
        offer.reward = request.POST.get('reward') or 0
        offer.activity_start = request.POST.get('activity_start')
        offer.activity_end = request.POST.get('activity_end')
        offer.landing_page = request.POST.get('landing_page') or 'https://partnetix.ru'
        offer.save()
    if (
        old_payout_type == 'fixed'
        and offer.payout_type == 'fixed'
        and old_reward != offer.reward
    ):
        paid_lead_ids = Accrual.objects.filter(
            lead__isnull=False,
            payout_status='paid'
        ).values_list('lead_id', flat=True)

        Lead.objects.filter(
            partner=partner,
            offer=offer
        ).exclude(
            id__in=paid_lead_ids
        ).update(
            partner_reward=offer.reward
        )
        delete_file_ids = request.POST.getlist('delete_promo_files')

        for file_id in delete_file_ids:
            if not file_id:
                continue

            promo_file = OfferPromoMaterial.objects.filter(
                id=file_id,
                offer=offer
            ).first()

            delete_file_ids = request.POST.getlist('delete_promo_files')

            for file_id in delete_file_ids:
                if not file_id:
                    continue

                promo_file = OfferPromoMaterial.objects.filter(
                    id=file_id,
                    offer=offer
                ).first()

                if promo_file:
                    file_name = promo_file.file.name if promo_file.file else ''

                    if file_name:
                        MarketingMaterial.objects.filter(
                            offer=offer,
                            file=file_name
                        ).delete()

                        MarketingMaterial.objects.filter(
                            offer=offer,
                            title=file_name.split('/')[-1]
                        ).delete()

                        promo_file.file.delete(save=False)

                    promo_file.delete()
        for uploaded_file in request.FILES.getlist('promo_materials'):

            promo_file = OfferPromoMaterial.objects.create(
                offer=offer,
                file=uploaded_file
            )

            MarketingMaterial.objects.create(
                title=uploaded_file.name,
                description='Промоматериал оффера',
                material_type=get_material_type_by_filename(
                    uploaded_file.name
                ),
                file=promo_file.file,
                offer=offer
            )

        messages.success(request, 'Оффер успешно обновлён.')

    return redirect('partner_detail', partner_id=partner.id)

def get_material_type_by_filename(filename):
    ext = filename.split('.')[-1].lower()

    image_exts = ['jpg', 'jpeg', 'png', 'webp', 'gif', 'svg']
    video_exts = ['mp4', 'mov', 'avi', 'webm', 'mkv']

    if ext in image_exts:
        return 'image'

    if ext in video_exts:
        return 'video'

    return 'text'

@login_required
@admin_required
def remove_partner_offer(request, partner_id, offer_id):

    partner = get_object_or_404(
        User,
        id=partner_id
    )

    partner_registration = get_object_or_404(
        PartnerRegistration,
        user=partner
    )

    offer = get_object_or_404(
        Offer,
        id=offer_id
    )

    # Удаляем связь партнёр ↔ оффер
    partner_registration.offers.remove(offer)

    # Если это основной оффер партнёра
    if partner_registration.offer_id == offer.id:
        partner_registration.offer = None
        partner_registration.save(update_fields=['offer'])

    # Удаляем все ссылки партнёра на этот оффер
    PartnerLink.objects.filter(
        partner=partner,
        offer=offer
    ).delete()

    messages.success(
        request,
        'Оффер отключён от партнёра.'
    )

    return redirect(
        'partner_detail',
        partner_id=partner.id
    )
def recalculate_partner_rewards(partner, partner_registration):
    paid_lead_ids = Accrual.objects.filter(
        lead__isnull=False
    ).filter(
        models.Q(payout_status='paid') |
        models.Q(withdrawal_requests__status='paid')
    ).values_list('lead_id', flat=True)

    leads_to_update = Lead.objects.filter(
        partner=partner,
        offer__payout_type='partner_status'
    ).exclude(
        id__in=paid_lead_ids
    ).select_related('tariff', 'offer')

    for lead in leads_to_update:
        if lead.status == 'deal' and lead.tariff:
            lead.deal_amount = lead.tariff.price
            lead.partner_reward = (
                lead.deal_amount
                * partner_registration.status.reward_percent
                / Decimal('100')
            )
        else:
            lead.partner_reward = Decimal('0')

        lead.save(update_fields=[
            'deal_amount',
            'partner_reward'
        ])

        Accrual.objects.filter(
            lead=lead
        ).update(
            amount=lead.partner_reward
        )